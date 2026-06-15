from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DOCX = Path("output/docx/final")
OUT_MD = Path("output/markdown")
OUT_DOCX.mkdir(parents=True, exist_ok=True)
OUT_MD.mkdir(parents=True, exist_ok=True)

TITLE = "Affärsmannaskap - levande instuderingsmaterial"
DOCX_PATH = OUT_DOCX / "Individuellt - Affärsmannaskap instuderingsmaterial.docx"
MD_PATH = OUT_MD / "Individuellt - Affärsmannaskap instuderingsmaterial.md"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_table_layout(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    return p


def write_cell(cell, text, bold=False, size=9.2):
    p = clear_cell(cell)
    parts = str(text).split("\n")
    for idx, part in enumerate(parts):
        if idx:
            p.add_run().add_break()
        r = p.add_run(part)
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.bold = bold


def new_decimal_num_id(doc):
    num = doc.part.numbering_part.element.add_num(7)
    num.add_lvlOverride(ilvl=0).add_startOverride(1)
    return num.numId


def apply_num(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
    return doc


def p(doc, text, style=None, bold=False):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(4 if style in ("List Bullet", "List Number") else 8)
    r = para.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.bold = bold
    return para


def title(doc, text, subtitle):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)
    r = para.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(26)
    r.font.bold = False
    r.font.color.rgb = RGBColor(0, 0, 0)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    sr = sub.add_run(subtitle)
    sr.font.name = "Arial"
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(85, 85, 85)


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def bullets(doc, items):
    for item in items:
        p(doc, item, "List Bullet")


def numbered(doc, items):
    num_id = new_decimal_num_id(doc)
    for item in items:
        para = p(doc, item, "List Number")
        apply_num(para, num_id)


def table(doc, headers, rows, widths, font_size=9.0):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_repeat_table_header(t.rows[0])
    set_row_cant_split(t.rows[0])
    for i, header in enumerate(headers):
        write_cell(t.rows[0].cells[i], header, bold=True, size=font_size)
        set_cell_shading(t.rows[0].cells[i], "F8F9FA")
    for row in rows:
        table_row = t.add_row()
        set_row_cant_split(table_row)
        cells = table_row.cells
        for i, value in enumerate(row):
            write_cell(cells[i], value, size=font_size)
    set_table_layout(t, widths)
    p(doc, "")
    return t


def md_table(headers, rows):
    out = ["| " + " | ".join(headers) + " |"]
    out.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        clean = [str(x).replace("\n", "<br>") for x in row]
        out.append("| " + " | ".join(clean) + " |")
    return "\n".join(out)


class MD:
    def __init__(self):
        self.lines = []

    def add(self, text=""):
        self.lines.append(text)

    def h1(self, text):
        self.add(f"# {text}")
        self.add()

    def h2(self, text):
        self.add(f"## {text}")
        self.add()

    def h3(self, text):
        self.add(f"### {text}")
        self.add()

    def p(self, text):
        self.add(text)
        self.add()

    def bullets(self, items):
        for item in items:
            self.add(f"- {item}")
        self.add()

    def numbered(self, items):
        for i, item in enumerate(items, 1):
            self.add(f"{i}. {item}")
        self.add()

    def table(self, headers, rows):
        self.add(md_table(headers, rows))
        self.add()


CONCEPTS = [
    ("Affärsmannaskap", "Att förstå kundens behov, skapa värde och fatta kommersiellt kloka beslut.", "Det visar om affären är hållbar, trovärdig och lönsam.", "Ett IT-bolag säljer inte bara automation, utan mindre administration och bättre kontroll.", "Att tro att affärsmannaskap bara betyder att sälja mer."),
    ("Konsultmässighet", "Att agera professionellt, kommunikativt och ansvarstagande i kundrelationer.", "Kunden köper både kompetens och förtroende.", "Konsulten återkopplar tidigt när en SharePoint-migrering riskerar att bli försenad.", "Att tro att teknisk kunskap räcker utan bemötande och tydlighet."),
    ("Kundvärde", "Den nytta kunden upplever jämfört med kostnad, risk och alternativ.", "Värde gör priset begripligt.", "Kunden sparar fem timmar administration per vecka genom automatiserade flöden.", "Att beskriva funktioner utan att koppla dem till nytta."),
    ("Värdeerbjudande", "En tydlig beskrivning av vilket problem ni löser och varför kunden ska välja er.", "Det är kärnan i erbjudandet.", "Vi hjälper små organisationer att automatisera återkommande administration i Microsoft 365.", "Att skriva en slogan i stället för ett konkret löfte."),
    ("Målgrupp/kundsegment", "Den grupp kunder ni riktar er till.", "Utan avgränsning blir erbjudandet för brett.", "Små och medelstora organisationer med mycket manuell ärendehantering.", "Att säga alla företag."),
    ("Positionering", "Hur ni vill uppfattas jämfört med alternativ.", "Det hjälper kunden förstå varför ni är ett rimligt val.", "Prisvärd verksamhetsnära Microsoft 365-partner för mindre organisationer.", "Att positionera sig som bäst på allt."),
    ("Differentiering", "Det som gör er mindre jämförbara med konkurrenterna.", "Differentiering minskar ren prisjämförelse.", "Ni kombinerar processförståelse, AI-stöd och praktisk SharePoint-implementation.", "Att påstå unikhet utan konkret skillnad."),
    ("Paketering", "Att samla tjänster i begripliga paket med tydligt innehåll.", "Det gör köpet enklare och kan förbättra marginalen.", "Förstudie, införandepaket och månadsvis förbättringsstöd.", "Att paketera så otydligt att kunden ändå måste gissa omfattningen."),
    ("Prissättning", "Hur priset sätts och motiveras.", "Pris visar värde, risk och prioritet.", "Fast pris för förstudie, projektpris för införande och abonnemang för stöd.", "Att använda rabatt som standardargument."),
    ("USP", "Det unika försäljningsargumentet.", "Det svarar på varför kunden ska välja just er.", "Vi bygger automation direkt i kundens befintliga Microsoft 365-miljö.", "Att kalla något unikt fast konkurrenter gör samma sak."),
    ("ESP", "Det känslomässiga försäljningsargumentet.", "Beslut påverkas också av trygghet, förtroende och enkelhet.", "Kunden känner att digitaliseringen blir hanterbar och inte ett stort IT-projekt.", "Att blanda ihop känsla med tom reklam."),
    ("FAB", "Features, Advantages, Benefits: egenskap, fördel, kundnytta.", "Modellen tvingar dig att gå från vad lösningen är till varför den spelar roll.", "Egenskap: Power Automate-flöde. Fördel: färre manuella steg. Nytta: snabbare handläggning.", "Att stanna vid egenskapsnivån."),
    ("5P", "Produkt, pris, plats, påverkan och personal.", "Ger en enkel helhetsbild av erbjudandet.", "Tjänst, prismodell, digital leverans, LinkedIn/referenser och konsulternas kompetens.", "Att bara prata produkt och glömma kanal, pris och människor."),
    ("SWOT", "Styrkor, svagheter, möjligheter och hot.", "Ger en överblick över intern och extern situation.", "Styrka: M365-kompetens. Hot: större konsulter och AI-verktyg som substitut.", "Att lägga externa hot som interna svagheter."),
    ("PESTELID", "Omvärldsanalys: politiskt, ekonomiskt, socialt, teknologiskt, legalt, internationellt, miljö och demografi.", "Visar faktorer utanför bolaget som påverkar affären.", "AI-regler, budgetläge, hybridarbete och kompetensbrist påverkar efterfrågan.", "Att fylla i allt utan att prioritera det viktiga."),
    ("Ansoff", "Matris för tillväxt med nuvarande/ny marknad och nuvarande/ny tjänst.", "Hjälper er se risknivå i tillväxtval.", "Sälja samma automationspaket till fler befintliga kunder inom små och medelstora organisationer är lägre risk än att skapa en ny egen mjukvarutjänst.", "Att tro att hög risk alltid är fel."),
    ("Business Model Canvas", "En översikt över hur affärsmodellen hänger ihop.", "Kopplar kund, erbjudande, resurser, kostnader och intäkter.", "Kundsegment, värdeerbjudande, kanaler, partners och intäktsströmmar för ett IT-konsultbolag.", "Att fylla rutorna utan röd tråd."),
    ("Produkt-marknadsmatris", "Ett sätt att poängsätta produkt- och marknadsalternativ.", "Hjälper gruppen prioritera realistiska affärer.", "AI-stöd för administration får hög kompetens- och lönsamhetspoäng i segmentet små och medelstora organisationer.", "Att låta magkänsla ersätta gemensam skala."),
    ("Konkurrentanalys", "Jämförelse av direkta, indirekta och substitutkonkurrenter.", "Visar vad kunden jämför er med.", "Stora konsulter, frilansare, interna IT-avdelningar och färdiga AI-verktyg.", "Att bara analysera företag som ser exakt likadana ut."),
    ("Företagsform", "Den juridiska formen för verksamheten.", "Påverkar ansvar, ägande, kapital, skatt och administration.", "Aktiebolag passar ofta ett växande konsultcase med flera ägare.", "Att välja form utan att diskutera ansvarsrisk."),
    ("Employer branding", "Hur arbetsgivaren uppfattas av nuvarande och framtida medarbetare.", "Konsultbolag behöver attrahera och behålla kompetens.", "Bolaget lyfter rimlig arbetsbelastning, utveckling och flexibelt arbete.", "Att tro att employer branding bara är marknadsföring."),
    ("Omvärldsanalys", "Att systematiskt förstå yttre faktorer som påverkar affären.", "Gör affärsargumenten mindre gissningsbaserade.", "REBR 2026 används för att resonera om arbetsliv, rekrytering och digital arbetsmiljö.", "Att använda externa källor som bevis för mer än de faktiskt visar."),
    ("Vision", "En inspirerande riktning för vart bolaget vill sträva.", "Den hjälper gruppen prioritera och förklara vart bolaget är på väg.", "Vi vill göra smart digitalisering möjlig för organisationer utan stor IT-avdelning.", "Att skriva ett kortsiktigt mål och kalla det vision."),
    ("Mission", "Bolagets varför och hur det bidrar till visionen.", "Den gör syftet tydligare än bara vad bolaget säljer.", "Vi förenklar administration med AI-stöd och Microsoft 365-lösningar som kunder faktiskt kan använda.", "Att mission blir samma sak som affärsidé."),
    ("Värderingar", "Det bolaget står för och de beteenden som ska prägla arbetet.", "Värderingar påverkar kultur, ledarskap och kundrelation.", "Tydlighet betyder att kunden alltid ska förstå status, risk och nästa steg.", "Att skriva fina ord som inte syns i beteenden."),
    ("Affärsidé", "Vad bolaget gör varje dag, för vem och vilken nytta det skapar.", "Den kopplar ihop erbjudande, kund och värde i vardagen.", "Vi kartlägger, automatiserar och förbättrar administrativa flöden för små och medelstora organisationer.", "Att affärsidén blir för bred och saknar kundnytta."),
    ("KPI", "Ett nyckeltal som visar om ett mål följs upp.", "Det gör mål mätbara och lättare att justera.", "Andel projekt levererade i tid, kundnöjdhet eller återkommande intäkter.", "Att mäta sådant som är lätt att mäta men inte viktigt."),
    ("Balanced Scorecard", "En modell för mål ur flera perspektiv: kund, medarbetare, interna processer och finansiellt.", "Den hindrar gruppen från att bara sätta ekonomiska mål.", "Kundnöjdhet, konsultutveckling, leveransprecision och marginal följs upp parallellt.", "Att använda modellen som tabell utan att koppla målen till affären."),
    ("Organisering", "Hur bolaget bemannar, leder och strukturerar arbetet runt sin idé.", "Organisationen måste passa erbjudandet, kulturen och målen.", "Ett konsultbolag kan behöva tydliga roller för sälj, leverans, kvalitet och kompetensutveckling.", "Att tro att organisering bara är ett organisationsschema."),
    ("Resultaträkning", "Visar intäkter minus kostnader under en period.", "Den visar om affären faktiskt går med vinst eller förlust.", "Essiq AB hade 2021 ett resultat efter finansiella poster på 11 305 793 kr.", "Att bara titta på omsättning och glömma kostnaderna."),
    ("Balansräkning", "Visar tillgångar, skulder och eget kapital vid en viss tidpunkt.", "Den visar hur bolaget är finansierat och vilken risk som finns i skulderna.", "Hälsö Fisk AB hade 1 184 841 kr i tillgångar och 130 528 kr i eget kapital 2017.", "Att läsa balansräkningen som om den visar försäljning över tid."),
    ("Eget kapital", "Tillgångar minus skulder; ägarnas kapital i bolaget.", "Det fungerar som buffert och påverkar soliditeten.", "Essiq AB hade 10 468 395 kr i eget kapital 2021.", "Att glömma skillnaden mellan bundet och fritt eget kapital."),
    ("Likviditet", "Kortsiktig betalningsförmåga.", "Ett bolag kan vara lönsamt men ändå få problem om pengar inte kommer in i tid.", "Essiq AB hade hög balanslikviditet men låg ren kassa jämfört med fordringarna.", "Att tro att vinst automatiskt betyder pengar på banken."),
    ("Kassalikviditet", "Omsättningstillgångar exklusive lager och pågående arbeten dividerat med kortfristiga skulder.", "Måttet visar om korta skulder kan betalas utan att vara beroende av lagerförsäljning.", "Hälsö Fisk AB låg ungefär runt 80 procent i kassalikviditet utifrån 2017 års balansräkning.", "Att räkna med lager när frågan gäller snabb betalningsförmåga."),
    ("Balanslikviditet", "Omsättningstillgångar dividerat med kortfristiga skulder.", "Ger en bredare bild av kortsiktig betalningsförmåga.", "Essiq AB hade cirka 144 procent i balanslikviditet 2021.", "Att inte se att fordringar måste betalas in för att bli kassa."),
    ("Soliditet", "Eget kapital i relation till totala tillgångar.", "Måttet visar långsiktig finansiell motståndskraft.", "Essiq AB rapporterade 28,20 procent soliditet 2021; Hälsö Fisk AB rapporterade 11 procent 2017.", "Att jämföra soliditet utan att förstå bransch, fas och obeskattade reserver."),
    ("Kassaflöde", "Pengar in och ut över tid.", "Kassaflödet avgör om bolaget kan betala löpande även när resultatet ser bra ut.", "Essiq AB:s likvida medel minskade från 26 509 757 kr till 2 822 630 kr under 2021.", "Att bara titta på årets resultat."),
    ("Täckningsbidrag", "Försäljningsintäkt minus direkta kostnader för en tjänst eller ett projekt.", "Visar hur mycket projektet bidrar till fasta kostnader och vinst.", "Ett projekt som säljs för 180 000 kr och kostar 105 000 kr direkt ger 75 000 kr i TB.", "Att kalla hela intäkten för vinst."),
    ("Budget och prognos", "Budget är plan, prognos är uppdaterad bedömning av hur det går.", "Det hjälper gruppen se avvikelser innan de blir akuta.", "Ett konsultbolag följer intäkt och beläggning per kvartal, tertial eller halvår.", "Att göra budget en gång och sedan aldrig justera."),
    ("Scenario", "Ett tänkbart utfall i budgeten, till exempel låg, normal eller hög försäljning.", "Scenarier gör budgeten mindre låst och hjälper gruppen se risk.", "Konsultbolaget räknar på vad som händer vid 50, 70 och 85 procent debiteringsgrad.", "Att bara räkna på ett optimistiskt läge."),
    ("Sociala avgifter", "Arbetsgivarens lagstadgade avgifter ovanpå bruttolönen.", "De gör att lönekostnaden är högre än den lön medarbetaren ser.", "En konsults månadskostnad måste räknas med lön, sociala avgifter, pension och andra personalkostnader.", "Att budgetera endast bruttolön."),
    ("Semesterlön/semestertillägg", "Ersättning för semester som behöver räknas in i personalkostnaden.", "Den påverkar verklig kostnad för anställda och därmed pris och beläggningskrav.", "Föreläsning 6 tar upp 12 procent vid 25 semesterdagar och justering med 0,48 procent per extra dag.", "Att glömma semesterkostnaden i konsultkalkylen."),
    ("Overhead (OH)", "Indirekta kostnader som behövs för att verksamheten ska fungera men inte hör till ett enskilt uppdrag.", "OH måste bäras av intäkterna från debiterbart arbete.", "Utbildning, administration, licenser, kontor, försäkringar, sälj och intern tid.", "Att bara räkna konsultens lön och glömma allt runt omkring."),
    ("Debiteringsgrad/faktureringsgrad", "Andelen arbetstid som kan faktureras kund.", "Ett centralt lönsamhetsmått i konsultbolag.", "Om en konsult har mycket intern tid behöver timpris eller paketpris täcka färre fakturerbara timmar.", "Att anta att all arbetstid kan faktureras."),
    ("Affärsmodell", "Hur bolaget skapar, levererar och tar betalt för värde.", "Den avgör vilka intäkter, kostnader och risker bolaget får.", "Ett IT-bolag kan kombinera konsulttimmar, fastprisprojekt, abonnemang och licenser.", "Att bara beskriva tjänsten och inte hur pengar faktiskt tjänas."),
    ("Freemium", "Grundversionen är gratis medan kunden betalar för premiumfunktioner.", "Kan skapa låg tröskel in men kräver tydlig väg till betalning.", "En digital tjänst låter små team testa gratis men tar betalt för avancerad automation.", "Att få många gratisanvändare utan betalningsmodell."),
    ("Prenumeration", "Kunden betalar återkommande, ofta per månad eller år.", "Ger återkommande intäkter men kräver löpande värde.", "Support, förvaltning, licenser eller förbättringsstöd säljs som månadsavtal.", "Att behandla prenumeration som engångsprojekt."),
    ("Fastprisprojekt", "Leverans säljs till ett bestämt pris för en definierad omfattning.", "Kan vara attraktivt för kunden men flyttar risk till leverantören.", "Ett införandepaket för SharePoint säljs för fast pris efter avgränsad förstudie.", "Att sälja fast pris utan tydlig omfattning och ändringshantering."),
    ("Löpande räkning", "Kunden betalar för faktisk nedlagd tid eller resursförbrukning.", "Minskar leverantörens omfattningsrisk men kräver förtroende och uppföljning.", "En senior konsult faktureras per timme under ett förbättringsarbete.", "Att missa att kunden kan uppleva svag kostnadskontroll."),
    ("Ramavtal", "Ett avtal som sätter ramarna för ett samarbete och sedan avropas mot.", "Vanligt i offentlig sektor och större organisationer.", "Ett konsultbolag kvalificerar sig som leverantör och får separata avrop för konkreta uppdrag.", "Att tro att ramavtal automatiskt betyder garanterad försäljning."),
    ("Avrop", "Beställning eller konkret uppdrag som görs inom ett ramavtal.", "Det är ofta avropet som anger den faktiska leveransen.", "Kunden avropar en förstudie enligt villkoren i ramavtalet.", "Att inte skilja mellan ramens villkor och det enskilda uppdraget."),
    ("Letter of Intent", "En avsiktsförklaring inför ett möjligt samarbete.", "Kan visa riktning innan slutligt avtal finns, men behöver läsas försiktigt.", "Parterna skriver en LOI om att utreda ett gemensamt produktprojekt.", "Att behandla en avsiktsförklaring som färdigt leveransavtal."),
    ("Konsultavtal", "Avtal som reglerar uppdrag, ansvar, ersättning, avtalstid, sekretess och andra villkor.", "Det styr både leverans och risk.", "Avtalet anger omfattning, rapportering, betalning, underleverantörer och ansvar.", "Att starta uppdrag utan att omfattning, ansvar och ersättning är tydligt."),
    ("NDA/sekretessavtal", "Avtal eller klausul som reglerar hur konfidentiell information får användas.", "IT-projekt kan ge insyn i kundens data, processer och affärshemligheter.", "Konsulten får inte sprida kundens systemritningar eller interna processdata.", "Att tro att sekretess bara gäller källkod."),
    ("SLA", "Service Level Agreement: överenskommelse om servicenivå, tillgänglighet, svarstid eller åtgärdstid.", "Gör drift- och supportförväntningar mätbara.", "Support svarar inom viss tid och kritiska incidenter prioriteras enligt avtal.", "Att lova servicenivåer utan kapacitet eller uppföljning."),
    ("MSA", "Master Service Agreement: huvudavtal som reglerar grundvillkor för flera uppdrag.", "Minskar behovet att förhandla allt från början varje gång.", "Ett konsultbolag har ett MSA med kunden och separata uppdragsbilagor för projekt.", "Att blanda ihop huvudavtal med konkret uppdragsbeskrivning."),
    ("GDPR/personuppgifter", "Regler för behandling av personuppgifter.", "IT-lösningar, drift, support och analys kan innebära personuppgiftsbehandling.", "Kunden och leverantören behöver reda ut roller, dataflöden och säkerhet.", "Att se GDPR som en juridisk bilaga utan påverkan på lösningen."),
    ("Offentlig upphandling", "När offentlig sektor köper varor eller tjänster genom en reglerad process.", "IT-konsulter kan behöva förstå krav, kriterier, tidslinje, offentlighet och överprövningsrisk innan de räknar på affären.", "Ett konsultbolag analyserar en kommunal upphandling och avgör om kraven, villkoren och utvärderingsmodellen passar bolaget.", "Att behandla offentlig upphandling som en vanlig privat offertdialog."),
    ("LOU", "Lag (2016:1145) om offentlig upphandling. I kursen används den som ram för hur offentliga köp styrs.", "Den påverkar hur krav får ställas, hur anbud jämförs och hur processen kan överprövas.", "En kommun måste beskriva uppdrag och utvärderingskriterier så att leverantörer kan lämna jämförbara anbud.", "Att plugga lagnumret utan att förstå vad det betyder för affären."),
    ("Upphandlande myndighet/offentligt styrt organ", "Aktörer som omfattas av upphandlingsregler, till exempel statliga och kommunala myndigheter samt vissa offentligt styrda organ.", "Det avgör om kunden måste följa offentlig upphandlingsprocess.", "En region, kommun eller offentligt styrt bolag kan behöva upphandla IT-tjänster enligt reglerade former.", "Att bara titta på kundens namn och missa vilken juridisk roll kunden har."),
    ("Tröskelvärde", "Beloppsgräns som påverkar vilka upphandlingsregler och processkrav som gäller.", "Tröskelvärden varierar med kundtyp och tjänst, så de ska kontrolleras i det konkreta fallet.", "Föreläsning 7 varnar för att tröskelvärden och direktupphandlingsgränser varierar beroende på tjänst, kund och tidpunkt.", "Att memorera ett gammalt belopp som om det alltid gäller."),
    ("Direktupphandling", "En enklare inköpsform som kan användas under vissa beloppsgränser eller förutsättningar.", "Kan ge snabbare affär, men kunden måste fortfarande följa relevanta regler och dokumentera processen.", "Föreläsningen tar upp direktupphandlingsgränsen som ett exempel och markerar att beloppet kan variera.", "Att tro att direktupphandling betyder helt fri beställning utan krav."),
    ("RFI/RFP/RFQ", "Tre olika förfrågningar: information, förslag/offert och pris på tydligt definierat behov.", "De visar hur mogen kundens köpprocess är och hur leverantören bör svara.", "RFI samlar marknadsinformation, RFP ber leverantörer föreslå lösning och RFQ används när kunden vet exakt vad som ska köpas.", "Att svara med pris när kunden egentligen söker information, eller tvärtom."),
    ("Avtalsspärr/överprövning", "En period efter tilldelning då avtal normalt inte får tecknas och leverantörer kan begära överprövning.", "Otydliga krav, orimliga krav eller otydlig värderingsmodell kan skapa processrisk.", "Föreläsning 7 nämner att överprövning ofta rör otydliga krav, orimliga krav eller oklar utvärdering.", "Att bara räkna på leverans och glömma risken att processen drar ut på tiden."),
    ("Offentlig handling och sekretess", "Material i offentlig upphandling kan bli offentlig handling, men vissa delar kan sekretessbeläggas.", "Leverantören måste skriva anbud så att affärshemligheter hanteras medvetet.", "Ett anbud kan behöva markera känslig information, men allt kan inte förutsättas bli hemligt.", "Att skriva hela anbudet som om ingen annan någonsin kan läsa det."),
    ("Affärsplan", "Ett samlat dokument som beskriver idé, kunder, mål, resurser, budget och aktiviteter.", "Den visar om affären hänger ihop från kundproblem till intäkter, leverans och uppföljning.", "Affärsplanen för gruppen bör visa differentiering, affärsmodell, SWOT och hur gruppen har arbetat fram riktning och mål.", "Att göra affärsplanen till en textlista utan beslut och röd tråd."),
    ("Roadmap", "En förenklad version av affärsplanen som visar riktning, prioriteringar och steg över tid.", "Den används ofta internt för att kommunicera vad som ska göras när.", "Roadmapen kan visa när bolaget ska göra förstudiepaket, första säljinsats, första pilotkund och återkommande supportmodell.", "Att blanda ihop roadmap med full affärsplan eller detaljerad projektplan."),
    ("Marknadsplan/marknadsbearbetning", "Plan för hur bolaget når, attraherar och bearbetar kunder.", "Kopplar målgrupp, kanal, budskap, säljaktivitet och prioritering till affären.", "LinkedIn, seminarier, partnerskap, referenser och direktkontakt kan prioriteras olika beroende på kundsegment.", "Att säga sociala medier utan att veta vilken kund som ska nås och varför."),
    ("Årsredovisning", "Paket med förvaltningsberättelse, resultaträkning, balansräkning, noter och ofta revisionsberättelse.", "Den ger flera vinklar på hur bolaget mår.", "AM5 använder Hälsö Fisk AB och Essiq AB som läsexempel.", "Att bara läsa en siffra utan att läsa sammanhanget."),
]


MODELS = [
    ("Business Model Canvas", "Samla affärsmodellen på en sida: kund, värde, leverans, resurser, kostnader och intäkter.", "När gruppen har en idé men behöver se om affären hänger ihop.", "Hur skapar, levererar och tar vi betalt för värde?", "Verksamhetslyftet Digital AB säljer en förstudie, ett införandepaket och löpande förbättringsstöd för Microsoft 365 och automation.", ["Rutorna fylls utan koppling till varandra.", "Intäkter och kostnader blir orealistiska.", "Kundsegmentet är för brett."], ["Kan vi beskriva kunden tydligt?", "Hänger intäktsströmmar ihop med erbjudandet?", "Vilka resurser måste finnas för att leverera?"]),
    ("SWOT", "Sammanfatta interna styrkor/svagheter och externa möjligheter/hot.", "Efter att ni har gjort en första marknads- och konkurrentbild.", "Vad talar för och emot vår affär?", "Styrka: M365 och processförståelse. Svaghet: få referenser. Möjlighet: behov av AI-stöd i administration. Hot: större konsultbolag och kunders interna IT.", ["Samma sak skrivs i flera rutor.", "Möjligheter blandas ihop med egna styrkor.", "SWOT används som åsiktslista utan beslut."], ["Vilka punkter är interna?", "Vilka punkter är externa?", "Vilket beslut följer av SWOT:en?"]),
    ("PESTELID", "Identifiera omvärldsfaktorer och bedöma påverkan, riktning och tidsram.", "När ni behöver visa att affären påverkas av omvärlden.", "Vilka yttre faktorer kan hjälpa eller skada vår affär?", "Teknologiskt: AI och automation ökar intresset. Legalt: dataskydd och AI-regler kräver kontroll. Ekonomiskt: kunder vill se tydlig nytta innan de köper.", ["För många faktorer utan prioritering.", "Allt blir hot eller allt blir möjlighet.", "Ingen koppling till det egna erbjudandet."], ["Vilka tre faktorer påverkar oss mest?", "Är påverkan positiv eller negativ?", "Vad behöver vi ändra i erbjudandet?"]),
    ("Ansoff", "Värdera tillväxtvägar genom ny/nuvarande tjänst och ny/nuvarande marknad.", "När gruppen diskuterar hur bolaget kan växa.", "Vilken typ av tillväxt väljer vi och hur riskfylld är den?", "Samma automationspaket till fler kunder inom små och medelstora organisationer är marknadspenetration. AI-stöd som ny tjänst till befintliga kunder är produktutveckling.", ["Alla idéer kallas expansion utan riskdiskussion.", "Ny marknad och ny tjänst underskattas.", "Risknivå används som absolut nej i stället för beslutsstöd."], ["Är tjänsten ny eller befintlig?", "Är marknaden ny eller befintlig?", "Vilken bevisning krävs för högre risk?"]),
    ("Produkt-marknadsmatris", "Poängsätta kombinationer av kundsegment och tjänster.", "När gruppen ska välja fokus i stället för att sälja allt till alla.", "Vilken kombination är mest attraktiv och realistisk?", "AI-stöd för administration till mindre organisationer får hög poäng på kompetens och kundbehov, medan avancerad egen mjukvarutjänst får högre risk.", ["Skalan 1-5 används olika av olika personer.", "Konkurrens bedöms inte ärligt.", "Kompetens väger för lite."], ["Vilken kombination får högst poäng?", "Varför är den realistisk?", "Vilken idé ska vänta?"]),
    ("Konkurrentanalys", "Förstå vilka alternativ kunden jämför med.", "När ni ska motivera positionering, pris och differentiering.", "Varför ska kunden välja oss i stället för alternativen?", "Alternativ kan vara stor konsultfirma, lokal IT-byrå, frilansare, intern IT eller färdiga AI-verktyg.", ["Bara direkta konkurrenter analyseras.", "Pris jämförs utan värde.", "Substitut glöms bort."], ["Vilka är direkta konkurrenter?", "Vilka substitut finns?", "Var är vår tydliga skillnad?"]),
    ("FAB", "Översätta egenskaper till fördelar och kundnytta.", "När ni ska skriva säljargument eller presentation.", "Vad betyder vår lösning för kunden?", "Egenskap: SharePoint-struktur. Fördel: enklare dokumentflöde. Kundnytta: mindre tid på att leta och rätt version används.", ["Argumentet stannar vid teknik.", "Nytta mäts inte.", "Fördel och benefit blandas ihop."], ["Vilken egenskap nämner vi?", "Vilken fördel ger den?", "Vilken nytta får kunden?"]),
    ("5P", "Se erbjudandet ur flera marknadsperspektiv.", "När ni ska förklara hela erbjudandet, inte bara tjänsten.", "Har vi tänkt på produkt, pris, plats, påverkan och personal?", "Produkt: automationspaket. Pris: fast förstudie och abonnemang. Plats: digital leverans. Påverkan: LinkedIn och referenser. Personal: konsulter med verksamhetsförståelse.", ["Bara produkten beskrivs.", "Personalens roll underskattas.", "Påverkan blir reklam utan målgrupp."], ["Vad säljer vi?", "Hur tar vi betalt?", "Hur får kunden reda på att vi finns?"]),
    ("Företagsplattform", "Samla vision, mission, värderingar och affärsidé så att bolaget står på en tydlig grund.", "När gruppen har valt affärsidé och behöver organisera sig runt den.", "Vilka är vi, varför finns vi och hur ska det synas i beteenden?", "Vision: smart digitalisering för mindre organisationer. Mission: minska administration med praktiska M365-lösningar. Värderingar: tydlighet, ansvar och lärande.", ["Vision, mission och affärsidé skrivs som nästan samma mening.", "Värderingar blir ord utan beteenden.", "Plattformen kopplas inte till kundnytta."], ["Är visionen framåtriktad?", "Beskriver missionen varför ni finns?", "Syns värderingarna i konkreta beteenden?"]),
    ("Mål och Balanced Scorecard", "Sätta mål och nyckeltal ur flera perspektiv.", "När gruppen ska visa hur bolaget följer upp framgång och inte bara beskriver erbjudandet.", "Hur vet vi att bolaget rör sig åt rätt håll?", "Kund: kundnöjdhet. Medarbetare: kompetensutveckling. Interna processer: leverans i tid. Finansiellt: marginal och återkommande intäkter.", ["Bara finansiella mål används.", "KPI:er saknar koppling till mål.", "Målen är för stora eller omöjliga att följa upp."], ["Vilka fyra perspektiv följer vi upp?", "Vilka nyckeltal visar verklig framdrift?", "Hur hanterar vi avvikelser?"]),
    ("Ekonomisk hälsokoll", "Läsa resultat, balansräkning, likviditet, soliditet och kassaflöde tillsammans.", "När gruppen ska visa om affären är ekonomiskt hållbar och inte bara snyggt beskriven.", "Mår bolaget bra nog för att fortsätta, växa och ta risk?", "Essiq AB visar positivt resultat och högre soliditet än Hälsö Fisk AB, men också kraftigt minskade likvida medel under 2021.", ["Omsättning används som bevis för lönsamhet.", "Likviditet och soliditet blandas ihop.", "En siffra tas ur årsredovisningen utan sammanhang."], ["Finns vinst eller förlust?", "Kan korta skulder betalas?", "Hur stark är kapitalbasen?", "Vad säger kassaflödet?"]),
    ("Årsredovisningsläsning", "Läsa förvaltningsberättelse, resultaträkning, balansräkning, kassaflöde, noter och revisionsberättelse i rätt ordning.", "När ni ska förstå varför två bolag med helt olika verksamheter ser olika ut ekonomiskt.", "Vilka delar av årsredovisningen stödjer samma berättelse om bolagets hälsa?", "Hälsö Fisk AB:s förvaltningsberättelse förklarar pressad försäljning med förändringar i Kungälv; Essiq AB beskriver hög efterfrågan och hög beläggning.", ["Resultaträkningen läses utan förvaltningsberättelsen.", "Jämförelseperioder jämförs rakt av trots olika längd.", "Rapporterad soliditet ersätts med egen förenklad kvot utan notering."], ["Vad säger bolaget själv har hänt?", "Vilka siffror stödjer berättelsen?", "Finns osäkerheter i perioden eller nyckeltalen?"]),
    ("Budget/prognos i konsultbolag", "Bygga en realistisk framåtblick över intäkter, kostnader, beläggning och scenarier.", "När gruppen ska visa om affären kan bära löner, pensioner, försäkringar, licenser, sälj och intern tid.", "Vilken intäkt krävs för att bolaget ska vara lönsamt i olika scenarier?", "Ett konsultbolag räknar på lön, sociala avgifter, pension, utbildning, licenser och debiteringsgrad innan pris sätts.", ["Blanda ihop budget och prognos.", "Glömma personalkringkostnader.", "Räkna med 100 procent fakturerbar tid."], ["Vilka kostnader finns även när inga nya projekt säljs?", "Vad händer om debiteringsgraden sjunker?", "När behöver prognosen revideras?"]),
    ("Avtalskarta för IT-projekt", "Identifiera vilka avtal, nivåer och riskfrågor som styr ett IT-uppdrag.", "När gruppen ska visa affärsmässig kontroll över leverans, ansvar, data, drift och ersättning.", "Vad måste vara reglerat innan kunden och leverantören bör starta?", "Ramavtal kan sätta ramen, avropet beskriver uppdraget och konsultavtalet reglerar omfattning, ansvar, ersättning och sekretess.", ["Tro att muntlig samsyn räcker.", "Glömma ägande, hosting, underhåll och personuppgifter.", "Lova SLA utan leveransförmåga."], ["Vem äger lösningen?", "Vem hostar och underhåller?", "Hur regleras ändringar, ansvar och data?"]),
    ("Upphandlingsanalys", "Läsa en offentlig upphandling som affärsbeslut: krav, villkor, volym, kriterier, risk och leveransförmåga.", "När kunden är offentlig eller när gruppen diskuterar hur ett konsultbolag kan sälja till kommun, region eller myndighet.", "Ska vi lämna anbud, avstå eller ställa frågor innan vi går vidare?", "I Mölndals Stad-övningen behöver gruppen reda ut vad kunden vill ha levererat, villkor, antal användare, tidigare data och vilka kriterier som avgör.", ["Svara innan kraven förståtts.", "Missa avtalsspärr och överprövningsrisk.", "Bara titta på pris och glömma kvalitet/värderingsmodell."], ["Vad ska levereras?", "Vilka krav är avgörande?", "Är anbuden jämförbara?", "Vilken risk tar vi om vi svarar?"]),
    ("Affärsplan och roadmap", "Samla affärsidé, kunder, marknad, konkurrens, produkt, försäljning, affärsmodell, resurser, hållbarhet, team och analys.", "När gruppen ska lämna in eller presentera bolagets samlade affärslogik.", "Hänger riktning, mål, analys, budget och aktiviteter ihop?", "Affärsplanen visar helheten, medan roadmapen visar de viktigaste stegen framåt på ett enklare sätt.", ["Lista rubriker utan beslut.", "Skriva SWOT/BMC utan att visa hur de påverkat riktningen.", "Göra roadmapen lika tung som affärsplanen."], ["Vilken riktning har analysen lett till?", "Vilka mål ska följas upp?", "Vilka aktiviteter gör planen verklig?"]),
]


COMPANY_FORMS = [
    ("Enskild näringsverksamhet", "En person driver verksamheten i eget namn.", "Solo-konsult eller liten testverksamhet.", "Enkelt att starta, inget kapitalkrav.", "Personligt ansvar för företagets skulder och svårare med flera ägare.", "Högre personlig risk.", "Passar dåligt för gruppens fiktiva växande IT-konsultbolag."),
    ("Aktiebolag", "En juridisk person som kan ägas av en eller flera aktieägare.", "Konsultbolag med flera grundare, tillväxtambition och behov av tydlig ansvarsskillnad.", "Trovärdigt, tydligare ägande, bolaget ansvarar normalt för skulderna.", "Kapitalkrav och mer administration.", "Lägre personlig risk än enskild firma och handelsbolag.", "Passar troligen bäst i grupparbetet."),
    ("Handelsbolag", "Minst två bolagsmän driver verksamhet tillsammans.", "Enklare samarbete mellan flera personer.", "Inget kapitalkrav och relativt lätt att förstå.", "Bolagsmännen är personligt och solidariskt ansvariga.", "Hög risk eftersom en delägare kan påverka andras ansvar.", "Passar sämre för ett konsultbolag med risk och flera ägare."),
    ("Kommanditbolag", "En variant av handelsbolag med minst en komplementär och en kommanditdelägare.", "När en part ska ha större ansvar och andra begränsat ansvar.", "Möjliggör olika ansvarsnivåer.", "Komplementären har personligt ansvar och formen är svårare att förklara.", "Blandad risk, hög för komplementären.", "Passar normalt inte bäst för ett enkelt gruppcase."),
    ("Ekonomisk förening", "Minst tre medlemmar driver verksamhet för medlemmarnas ekonomiska nytta.", "Medlemsägd samverkan, plattform eller specialistpool.", "Demokratisk struktur och tydligt medlemsnyttoperspektiv.", "Mindre naturlig för ett vanligt konsultbolag med ägare och kunder.", "Mellanrisk beroende på upplägg.", "Passar bara om caset bygger på medlemsnytta."),
]


def add_concept_section(doc, md):
    h1(doc, "3. Begrepp jag måste kunna")
    md.h2("3. Begrepp jag måste kunna")
    p(doc, "Läs begreppen aktivt: förklara dem högt, koppla dem till ett exempel och kontrollera att du inte blandar ihop närliggande ord.")
    md.p("Läs begreppen aktivt: förklara dem högt, koppla dem till ett exempel och kontrollera att du inte blandar ihop närliggande ord.")
    headers = ["Begrepp", "Förstå", "Tillämpa och akta dig för"]
    rows = [
        (
            concept,
            f"Enkel förklaring: {explanation}\nVarför viktigt: {why}",
            f"Exempel: {example}\nRisk för missförstånd: {risk}",
        )
        for concept, explanation, why, example, risk in CONCEPTS
    ]
    table(doc, headers, rows, [1.35, 2.55, 2.6], font_size=8.0)
    md.table(headers, rows)


def add_models_section(doc, md):
    h1(doc, "4. Modellerna - vad de gör och när de används")
    md.h2("4. Modellerna - vad de gör och när de används")
    p(doc, "Modellerna ska hjälpa dig fatta och förklara beslut. De är inte mål i sig. För varje modell bör du kunna säga vilken fråga den svarar på och vilket beslut den leder till.")
    md.p("Modellerna ska hjälpa dig fatta och förklara beslut. De är inte mål i sig. För varje modell bör du kunna säga vilken fråga den svarar på och vilket beslut den leder till.")
    for name, use, when, question, example, mistakes, controls in MODELS:
        h2(doc, name)
        md.h3(name)
        rows = [
            ("Vad modellen används till", use),
            ("När i grupparbetet den används", when),
            ("Vilken fråga modellen hjälper oss att svara på", question),
            ("Mini-exempel", f"Pedagogiskt exempel: {example}"),
            ("Vanliga fel", "; ".join(mistakes)),
            ("Kontrollfrågor", "; ".join(controls)),
        ]
        table(doc, ["Del", "Innehåll"], rows, [2.05, 4.45], font_size=8.5)
        md.table(["Del", "Innehåll"], rows)


def add_annual_reports_section(doc, md):
    h1(doc, "8. Miniuppgift: använd AM5 på två årsredovisningar")
    md.h2("8. Miniuppgift: använd AM5 på två årsredovisningar")
    p(doc, "Det här är inte en uppgift som du ska besvara. Det är ett färdigt arbetat exempel som visar hur AM5-begreppen kan användas för att läsa två helt olika bolag: Hälsö Fisk AB, org.nr 556723-8497, och Essiq AB, org.nr 556674-6250.")
    md.p("Det här är inte en uppgift som du ska besvara. Det är ett färdigt arbetat exempel som visar hur AM5-begreppen kan användas för att läsa två helt olika bolag: Hälsö Fisk AB, org.nr 556723-8497, och Essiq AB, org.nr 556674-6250.")
    h2(doc, "Siffror som används i exemplet")
    md.h3("Siffror som används i exemplet")
    figure_rows = [
        ("Nettoomsättning", "13 368 849 kr, 2017", "184 605 717 kr, 2021"),
        ("Resultat efter finansiella poster", "-44 903 kr", "11 305 793 kr"),
        ("Årets resultat", "-44 903 kr", "7 636 962 kr"),
        ("Summa tillgångar", "1 184 841 kr", "72 634 351 kr"),
        ("Eget kapital", "130 528 kr", "10 468 395 kr"),
        ("Kortfristiga skulder", "1 054 313 kr", "49 244 570 kr"),
        ("Kassa och bank", "332 542 kr", "2 822 630 kr"),
        ("Rapporterad soliditet", "11 procent", "28,20 procent"),
    ]
    table(doc, ["Nyckelpost", "Hälsö Fisk AB", "Essiq AB"], figure_rows, [1.7, 2.15, 2.15], font_size=8.2)
    md.table(["Nyckelpost", "Hälsö Fisk AB", "Essiq AB"], figure_rows)
    p(doc, "Källnot: Hälsö Fisk AB:s dokument gäller räkenskapsåret 2017-01-01 till 2017-12-31. Essiq AB:s årsredovisning gäller 2021, men jämförelsekolumnen i resultaträkningen är perioden 2019-07-01 till 2020-12-31, alltså 18 månader. Jämför därför inte perioderna rakt av.")
    md.p("Källnot: Hälsö Fisk AB:s dokument gäller räkenskapsåret 2017-01-01 till 2017-12-31. Essiq AB:s årsredovisning gäller 2021, men jämförelsekolumnen i resultaträkningen är perioden 2019-07-01 till 2020-12-31, alltså 18 månader. Jämför därför inte perioderna rakt av.")
    qa_rows = [
        (
            "Vilket bolag går bäst resultatmässigt?",
            "Essiq AB ser klart starkare ut. Bolaget har 11 305 793 kr i resultat efter finansiella poster och 7 636 962 kr i årets resultat. Hälsö Fisk AB har -44 903 kr i resultat efter finansiella poster och årets resultat. AM5-poängen är att nettoomsättning inte räcker: kostnader och finansiella poster avgör om det blir vinst.",
        ),
        (
            "Vilket bolag har starkast soliditet?",
            "Essiq AB har högre rapporterad soliditet: 28,20 procent jämfört med Hälsö Fisk AB:s 11 procent. Det betyder att Essiq verkar ha bättre långsiktig motståndskraft. AM5 anger 30-40 procent som ungefärlig stark nivå, så Essiq ligger nära men under, medan Hälsö Fisk är tydligt svagare.",
        ),
        (
            "Hur ser den kortsiktiga betalningsförmågan ut?",
            "Hälsö Fisk AB är pressat: omsättningstillgångar på cirka 927 618 kr mot kortfristiga skulder på 1 054 313 kr ger ungefär 88 procent i balanslikviditet. Essiq AB har 70 915 128 kr i omsättningstillgångar mot 49 244 570 kr i kortfristiga skulder, cirka 144 procent. Men Essiqs rena kassa är bara 2 822 630 kr, så mycket ligger i fordringar.",
        ),
        (
            "Vad säger eget kapital om risk?",
            "Hälsö Fisk AB har 130 528 kr i eget kapital mot 1 184 841 kr i tillgångar. Det är en tunn buffert. Essiq AB har 10 468 395 kr i eget kapital och dessutom obeskattade reserver, vilket förklarar varför den rapporterade soliditeten blir högre än en enkel kvot eget kapital/tillgångar.",
        ),
        (
            "Vad säger förvaltningsberättelsen?",
            "Hälsö Fisk AB beskriver minskad försäljning och färre kundbesök på grund av infrastrukturförändringar i Kungälv. Det förklarar varför siffrorna ser pressade ut. Essiq AB beskriver hög efterfrågan, hög beläggningsgrad och rekrytering. Därför ser Essiqs problem mer ut som tillväxt- och kassaflödesstyrning än överlevnadsproblem.",
        ),
        (
            "Vad är den viktigaste lärdomen?",
            "Läs alltid flera delar tillsammans. Resultaträkningen visar om bolaget tjänar pengar. Balansräkningen visar kapital och skulder. Likviditet visar kort betalningsförmåga. Soliditet visar långsiktig motståndskraft. Förvaltningsberättelsen förklarar varför siffrorna ser ut som de gör.",
        ),
    ]
    h2(doc, "Frågor med färdiga svar")
    md.h3("Frågor med färdiga svar")
    for question, answer in qa_rows:
        h3(doc, question)
        md.h3(question)
        p(doc, answer)
        md.p(answer)


def add_lecture6_section(doc, md):
    h1(doc, "9. Föreläsning 6: budget, lönsamhet, affärsmodell och avtal")
    md.h2("9. Föreläsning 6: budget, lönsamhet, affärsmodell och avtal")
    p(doc, "Källa: Affärsmannaskap 6, Föreläsning 6, 2026-06-05. Föreläsningen fortsätter företagsekonomin men flyttar fokus från att läsa årsredovisningar till att själv styra en konsultaffär: budget, prognos, kostnader för anställda, lönsamhetslogik, affärsmodeller och avtal.")
    md.p("Källa: Affärsmannaskap 6, Föreläsning 6, 2026-06-05. Föreläsningen fortsätter företagsekonomin men flyttar fokus från att läsa årsredovisningar till att själv styra en konsultaffär: budget, prognos, kostnader för anställda, lönsamhetslogik, affärsmodeller och avtal.")
    p(doc, "Den praktiska poängen för en IT-projektledare är att tekniska beslut snabbt blir affärsbeslut: scope, bemanning, licenser, drift, ansvar, dataskydd och leveransform påverkar både kundvärde och lönsamhet.")
    md.p("Den praktiska poängen för en IT-projektledare är att tekniska beslut snabbt blir affärsbeslut: scope, bemanning, licenser, drift, ansvar, dataskydd och leveransform påverkar både kundvärde och lönsamhet.")

    h2(doc, "Budget och prognos")
    md.h3("Budget och prognos")
    p(doc, "Föreläsningen beskriver budget som en kvalificerad gissning om framtiden. En budget behöver därför visa både intäkter, kostnader och scenarier. Prognos/forecast är uppföljning och eventuell revidering av budgeten när verkligheten ändras.")
    md.p("Föreläsningen beskriver budget som en kvalificerad gissning om framtiden. En budget behöver därför visa både intäkter, kostnader och scenarier. Prognos/forecast är uppföljning och eventuell revidering av budgeten när verkligheten ändras.")
    budget_rows = [
        ("Kostnader", "Lön, sociala avgifter, pensioner, köpta tjänster, utrustning, telefon, licenser, lokal/hyra, resor, utbildning, försäkringar, kontorsmaterial, marknadsföring/sälj samt representation, möten och events."),
        ("Intäkter", "Arvoden, kick-back, royalty, licensintäkter och andra betalningsströmmar som passar affärsmodellen."),
        ("Scenarier", "Räkna inte bara på bästa läget. Gör exempelvis låg, normal och hög beläggning/försäljning så gruppen ser risk och känslighet."),
        ("Prognos", "Följ upp budgeten och revidera när antaganden om försäljning, kostnad, bemanning eller leverans inte längre stämmer."),
    ]
    table(doc, ["Del", "Vad gruppen ska få med"], budget_rows, [1.35, 5.15], font_size=8.4)
    md.table(["Del", "Vad gruppen ska få med"], budget_rows)

    h2(doc, "Vad kostar en anställd?")
    md.h3("Vad kostar en anställd?")
    employee_rows = [
        ("Synlig lön", "Bruttolönen är bara startpunkten. Arbetsgivaren måste också bära arbetsgivaravgifter, pension, semester och andra personalkostnader."),
        ("Försäkringar", "Föreläsningen tar upp företagsförsäkring och avtalsförsäkringar kopplade till exempelvis sjukdom, arbetsskada, arbetsbrist, dödsfall och föräldraledighet."),
        ("Pension", "Pensionskostnader behöver budgeteras separat. Föreläsningen anger ITP-exempel från 1 500 kr per månad, vilket ska läsas som kursens exempel, inte som universellt belopp."),
        ("Semester", "Semesterlönen tas upp som 12 procent vid 25 dagar och justering med 0,48 procent per extra semesterdag. Rörliga delar kan påverka."),
        ("Overhead", "Utbildning, administration, licenser, intern tid, sälj och annat som behövs för verksamheten måste täckas av debiterbara intäkter."),
    ]
    table(doc, ["Kostnadsdel", "Studiestöd"], employee_rows, [1.35, 5.15], font_size=8.3)
    md.table(["Kostnadsdel", "Studiestöd"], employee_rows)

    h2(doc, "Lönsamhet i konsultbolag")
    md.h3("Lönsamhet i konsultbolag")
    profitability_rows = [
        ("Tjänstemix", "Paketerade tjänster/produkter, fastprisprojekt, löpande räkning och fasta avtal ger olika risk, marginal och intäktsstabilitet."),
        ("Debiteringsgrad", "Ju lägre andel tid som faktureras kund, desto mer måste varje debiterbar timme eller varje paket bära."),
        ("Medarbetarkostnad", "Räkna med lön, sociala avgifter, pension, försäkring, utbildning, licenser och overhead innan priset bedöms."),
        ("Påverkbara faktorer", "Gruppen kan påverka paketering, pris, scope, beläggning, återkommande avtal, leveranseffektivitet och hur mycket intern tid som krävs."),
    ]
    table(doc, ["Fråga", "Hur den påverkar affären"], profitability_rows, [1.45, 5.05], font_size=8.5)
    md.table(["Fråga", "Hur den påverkar affären"], profitability_rows)

    h2(doc, "Affärsmodeller som togs upp")
    md.h3("Affärsmodeller som togs upp")
    business_model_rows = [
        ("Freemium", "Gratis grundversion, betalda premiumfunktioner. Exempel i föreläsningen: Spotify och Dropbox."),
        ("Prenumeration", "Återkommande avgift per månad eller år. Kan passa support, leasing, licenser och löpande förbättringsstöd."),
        ("Återförsäljare/e-handel", "Köper in och säljer vidare, fysiskt eller online. Exempel i föreläsningen: Amazon och Temu."),
        ("Franchise", "Varumärke och koncept licensieras ut mot avgift. Exempel i föreläsningen: McDonald's och Espresso House."),
        ("Cirkulär ekonomi", "Tjänar pengar på återanvändning, reparation eller uthyrning i stället för bara nyförsäljning."),
        ("P2P/plattform", "Mellanhand som kopplar ihop utbud och efterfrågan och tar avgift eller provision. Exempel i föreläsningen: Airbnb och Uber."),
        ("Rakhyvel och blad", "Grundprodukt säljs billigt medan tillbehör eller förbrukning säljs dyrare, till exempel skrivare och bläckpatroner."),
        ("Paketering", "Färdiga paketerbjudanden där kunden köper ett samlat innehåll, exempelvis ett Office-paket eller ett konsultpaket."),
    ]
    table(doc, ["Modell", "Enkel förklaring"], business_model_rows, [1.65, 4.85], font_size=8.2)
    md.table(["Modell", "Enkel förklaring"], business_model_rows)

    h2(doc, "Avtalsrätt och IT-projekt")
    md.h3("Avtalsrätt och IT-projekt")
    p(doc, "Föreläsningen använder avtalsrätt för att visa varför affären behöver tydliga ramar. Handslag och muntlig samsyn kan vara otillräckligt när uppdraget rör programvara, drift, data, ersättning, ansvar och ändringar.")
    md.p("Föreläsningen använder avtalsrätt för att visa varför affären behöver tydliga ramar. Handslag och muntlig samsyn kan vara otillräckligt när uppdraget rör programvara, drift, data, ersättning, ansvar och ändringar.")
    contract_level_rows = [
        ("Ramavtal", "Reglerar ramarna för samarbetet och avropas mot. Vanligt i offentlig sektor och större organisationer."),
        ("Avtal", "Reglerar ett mer specifikt tillfälle, en specifik leverans eller en specifik produkt/tjänst."),
        ("Letter of Intent", "Avsiktsförklaring inför ett möjligt samarbete. Ska inte förväxlas med ett färdigt leveransavtal."),
    ]
    table(doc, ["Nivå", "Studiestöd"], contract_level_rows, [1.55, 4.95], font_size=8.5)
    md.table(["Nivå", "Studiestöd"], contract_level_rows)

    h2(doc, "Kundens funderingar inför avtal")
    md.h3("Kundens funderingar inför avtal")
    contract_questions = [
        "Vem äger själva produkten eller programvaran?",
        "Vem hostar IT-systemet och hur hanteras drift och underhåll?",
        "Finns resurser att skala upp eller ned systemet vid behov?",
        "Vilka krav finns på systemets design och arkitektur?",
        "Hur hanteras open source?",
        "Hur regleras ersättning, underleverantörer och ansvar?",
        "Hur hanteras dataskydd, personuppgifter, sekretess, tvister och ansvarsbegränsningar?",
        "Agerar parterna som självständiga parter och är det tydligt vem som ansvarar för vad?",
    ]
    bullets(doc, contract_questions)
    md.bullets(contract_questions)

    h2(doc, "Konsultavtal och standardavtal")
    md.h3("Konsultavtal och standardavtal")
    consulting_rows = [
        ("Konsultavtal", "Uppdrag/omfattning, rapportering, konsultens och uppdragsgivarens ansvar, betalning/ersättning, avtalstid, sekretess/NDA, underleverantörer, konkurrens och överlåtelse."),
        ("IT och Telekomföretagens avtal", "Föreläsningen nämner Avtal 90, IT-Projekt, Agila Projekt, IT-Tjänster, IT-Infrastrukturtjänster och IT-Underhåll som exempel på standardavtal."),
        ("Fler avtalstyper", "Molntjänster, sekretessavtal, SLA, MSA, franchise, licens, royalties, patent, upphovsrätt och GDPR/personuppgifter."),
    ]
    table(doc, ["Område", "Vad som ska kommas ihåg"], consulting_rows, [1.6, 4.9], font_size=8.2)
    md.table(["Område", "Vad som ska kommas ihåg"], consulting_rows)


def add_lecture7_section(doc, md):
    h1(doc, "10. Föreläsning 7: offentlig upphandling, affärsplan och roadmap")
    md.h2("10. Föreläsning 7: offentlig upphandling, affärsplan och roadmap")
    p(doc, "Källor: Affärsmannaskap 7, Föreläsning 7, 2026-06-08 samt mallen ny-affarsplan-2026. Föreläsningen lägger till två praktiska områden: hur offentlig upphandling fungerar som köpprocess och hur gruppens affärsplan ska visa riktning, mål, analys, budget och aktiviteter.")
    md.p("Källor: Affärsmannaskap 7, Föreläsning 7, 2026-06-08 samt mallen ny-affarsplan-2026. Föreläsningen lägger till två praktiska områden: hur offentlig upphandling fungerar som köpprocess och hur gruppens affärsplan ska visa riktning, mål, analys, budget och aktiviteter.")
    p(doc, "Tentapoängen är att kunna koppla ihop upphandling med affärsmannaskap. En offentlig kund köper inte alltid genom fri dialog. Kunden styrs av regler, krav, kriterier, offentlighet och möjlighet till överprövning. För leverantören blir frågan därför: kan vi förstå uppdraget, uppfylla kraven och vinna på rätt grunder?")
    md.p("Tentapoängen är att kunna koppla ihop upphandling med affärsmannaskap. En offentlig kund köper inte alltid genom fri dialog. Kunden styrs av regler, krav, kriterier, offentlighet och möjlighet till överprövning. För leverantören blir frågan därför: kan vi förstå uppdraget, uppfylla kraven och vinna på rätt grunder?")

    h2(doc, "Offentlig upphandling i korthet")
    md.h3("Offentlig upphandling i korthet")
    procurement_rows = [
        ("LOU och reglerad process", "Föreläsningen tar upp LOU 2016:1145 och förtydliganden kring offentligt styrt organ och upphandlande myndighet. Poängen är att offentliga köp styrs av formella regler, inte bara av relation och förhandling."),
        ("Upphandlande myndighet/offentligt styrt organ", "Statliga och kommunala myndigheter samt vissa offentligt styrda organ kan omfattas. Det avgör om kunden måste köpa enligt upphandlingsregler."),
        ("Tröskelvärde och direktupphandling", "Tröskelvärden och direktupphandlingsgränser varierar med typ av kund, tjänst och tidpunkt. Föreläsningens belopp ska därför läsas som kurskontext, inte som evig regel."),
        ("Kriterier", "Kunden kan utvärdera bästa förhållande mellan pris och kvalitet, kostnad/livscykelkostnad eller pris. Som leverantör måste du förstå vad som faktiskt ger poäng."),
        ("Offentlig handling", "Anbud och upphandlingsmaterial kan bli offentliga handlingar. Vissa delar kan sekretessbeläggas, men sekretess ska inte tas för given."),
        ("Överprövning", "Föreläsningen beskriver överprövning inom avtalsspärren, ofta kopplad till otydliga krav, orimliga krav eller otydlig värderingsmodell."),
        ("Var upphandlingar hittas", "Föreläsningen nämner Opic och TendSign som platser där offentliga upphandlingar finns, samt Upphandling24 som informationskälla."),
    ]
    table(doc, ["Del", "Vad du ska kunna säga på tentan"], procurement_rows, [1.65, 4.85], font_size=8.2)
    md.table(["Del", "Vad du ska kunna säga på tentan"], procurement_rows)

    h2(doc, "RFI, RFP och RFQ")
    md.h3("RFI, RFP och RFQ")
    request_rows = [
        ("RFI", "Request for Information. Kunden samlar information om marknad, leverantörer eller tekniska lösningar inför ett större projekt eller upphandling. Svaret ska hjälpa kunden förstå möjligheter, inte bara ge pris."),
        ("RFP", "Request for Proposal/offertförfrågan. Kunden bjuder in leverantörer att lämna förslag och offert på produkt, tjänst eller projekt. Svaret behöver visa lösning, värde, upplägg och prislogik."),
        ("RFQ", "Request for Quotation. Kunden vet mer exakt vad som ska köpas och vill ha pris och villkor. Här blir jämförbarhet, leveranstid och pris tydligare."),
    ]
    table(doc, ["Förfrågan", "Hur du känner igen den"], request_rows, [1.1, 5.4], font_size=8.4)
    md.table(["Förfrågan", "Hur du känner igen den"], request_rows)

    h2(doc, "Så läser du en offentlig upphandling")
    md.h3("Så läser du en offentlig upphandling")
    reading_questions = [
        "Vad vill kunden ha levererat, konkret och avgränsat?",
        "På vilka villkor ska leveransen ske: tid, ansvar, avtal, ersättning, data, drift och support?",
        "För hur många användare, enheter, verksamheter eller processer gäller uppdraget?",
        "Vad ska hända med tidigare data, tidigare system eller befintliga arbetssätt?",
        "Vilka krav är skallkrav och vilka ger poäng?",
        "Vilka kriterier avgör vinnaren: pris, kostnad, kvalitet eller kombination?",
        "Är kraven tydliga nog för jämförbara anbud, eller finns risk för oklar värdering?",
        "Vilka delar av anbudet kan bli offentliga och vilka delar behöver sekretessmarkeras?",
        "Är uppdraget lönsamt med realistisk bemanning, risk och tidslinje?",
    ]
    bullets(doc, reading_questions)
    md.bullets(reading_questions)

    h2(doc, "Mölndals Stad-övningen")
    md.h3("Mölndals Stad-övningen")
    p(doc, "Föreläsningen använder en övning om Mölndals Stad. Den viktigaste studietekniken är att läsa upphandlingen som en affärs- och riskfråga, inte som en text som bara ska sammanfattas.")
    md.p("Föreläsningen använder en övning om Mölndals Stad. Den viktigaste studietekniken är att läsa upphandlingen som en affärs- och riskfråga, inte som en text som bara ska sammanfattas.")
    molndal_rows = [
        ("Vad vill de ha levererat?", "Identifiera faktisk leverans, omfattning och gränser. Om leveransen är otydlig blir både pris och risk svåra att bedöma."),
        ("På vilka villkor?", "Läs avtal, krav, ansvar, ersättning, support, tidsplan och eventuella sanktioner."),
        ("För hur många användare?", "Volym påverkar licenser, support, utbildning, drift och prissättning."),
        ("Vad händer med tidigare data?", "Migrering, rensning, ägande, säkerhet och ansvar kan bli större än själva nyinförandet."),
        ("Vilka kriterier avgör?", "Skilj mellan krav som måste uppfyllas och kriterier som ger konkurrensfördel."),
        ("Reflektion runt punkt 6.3", "Läs punkten som en risk- eller tolkningsfråga: vad kräver den, hur bevisar leverantören uppfyllelse och kan olika anbud jämföras rättvist?"),
    ]
    table(doc, ["Fråga", "Hur du pluggar på den"], molndal_rows, [1.75, 4.75], font_size=8.2)
    md.table(["Fråga", "Hur du pluggar på den"], molndal_rows)

    h2(doc, "Affärsplan, marknadsplan och roadmap")
    md.h3("Affärsplan, marknadsplan och roadmap")
    p(doc, "Föreläsningen säger att affärsplanen beskriver idé, kunder, mål, resurser, budget och aktiviteter för att nå målen. Den ska också visa differentiering, affärsmodell och SWOT, alltså hur gruppen arbetat fram riktning och mål. Utöver det ska en marknadsplan finnas med.")
    md.p("Föreläsningen säger att affärsplanen beskriver idé, kunder, mål, resurser, budget och aktiviteter för att nå målen. Den ska också visa differentiering, affärsmodell och SWOT, alltså hur gruppen arbetat fram riktning och mål. Utöver det ska en marknadsplan finnas med.")
    plan_rows = [
        ("1. Affärsidé/verksamhet", "Vad erbjuder ni och vilket problem eller behov löser ni hos kunden?"),
        ("2. Vision", "Vilken långsiktig ambition, position och vilket kundvärde vill företaget skapa?"),
        ("3. Mål", "Vilka konkreta mål finns på 3-5 års sikt: tillväxt, lönsamhet, kundbas, marknad eller organisation?"),
        ("4. Kunder", "Vilka målgrupper och segment har ni, vilka behov har de och varför väljer de er?"),
        ("5. Marknad", "Hur är marknaden avgränsad och vilka trender, möjligheter eller förändringar påverkar den?"),
        ("6. Konkurrens", "Vilka konkurrenter och alternativ finns, vilka styrkor/svagheter har de och hur skiljer ni er?"),
        ("7. Produkt/tjänst", "Vad är det konkreta värdet för kunden och vad gör erbjudandet särskiljande?"),
        ("8. Sälja, betala, leverera", "Hur når ni kunden, hur sker försäljningen, hur betalar kunden och hur levereras tjänsten?"),
        ("9. Marknadsbearbetning", "Vilka kanaler och aktiviteter använder ni, och hur prioriterar ni mellan dem?"),
        ("10. Affärsmodell", "Hur tjänar företaget pengar: prissättning, intäktsströmmar, kostnadsstruktur och marginaler?"),
        ("11. Resurser/förutsättningar", "Vilken kompetens, teknik, systemstöd, partners eller andra resurser krävs?"),
        ("12. Hållbart företagande", "Hur skapas långsiktigt ekonomiskt, miljömässigt och socialt värde?"),
        ("13. Teamet", "Vilka roller och kompetenser finns och hur kompletterar teamet varandra?"),
        ("14. Analys", "Sammanfatta styrkor, svagheter, möjligheter och hot, gärna kopplat till SWOT och PESTELID."),
    ]
    table(doc, ["Affärsplanens del", "Tentafrågan bakom rubriken"], plan_rows, [1.9, 4.6], font_size=7.7)
    md.table(["Affärsplanens del", "Tentafrågan bakom rubriken"], plan_rows)

    roadmap_rows = [
        ("Affärsplan", "Full helhet: kund, marknad, erbjudande, analys, ekonomi, resurser, team, hållbarhet och aktiviteter."),
        ("Marknadsplan", "Hur marknaden ska bearbetas: målgrupp, kanal, budskap, aktivitet och prioritering."),
        ("Roadmap", "Förenklad version av affärsplanen, ofta för intern kommunikation. Visar större steg, prioriteringar och tidsordning."),
        ("Scenarier", "Tänkbara utfall, till exempel låg/normal/hög försäljning eller olika kundvägar. Används för att visa risk och handlingsberedskap."),
    ]
    table(doc, ["Dokument/verktyg", "Skillnad du ska kunna förklara"], roadmap_rows, [1.55, 4.95], font_size=8.4)
    md.table(["Dokument/verktyg", "Skillnad du ska kunna förklara"], roadmap_rows)

    h2(doc, "Vanliga tentafällor från föreläsning 7")
    md.h3("Vanliga tentafällor från föreläsning 7")
    pitfalls = [
        "Att memorera en direktupphandlingsgräns som om den alltid gäller. Förklara hellre att gränser varierar och måste kontrolleras.",
        "Att blanda ihop RFI, RFP och RFQ. De visar olika grad av mognad i kundens behov.",
        "Att tro att lägsta pris alltid vinner. Kriterierna kan också handla om bästa förhållande mellan pris och kvalitet eller kostnad över livscykeln.",
        "Att glömma offentlig handling och sekretess när anbud skrivs.",
        "Att tro att en upphandling alltid leder till avtal. Föreläsningens BRG-exempel visar att processer kan dras tillbaka, göras om eller förskjutas.",
        "Att skriva affärsplanens rubriker utan att visa hur differentiering, affärsmodell, SWOT och marknadsplan leder till samma riktning.",
        "Att göra roadmapen till en full affärsplan i stället för en förenklad kommunikationsbild.",
    ]
    bullets(doc, pitfalls)
    md.bullets(pitfalls)


def build_document():
    doc = make_doc()
    md = MD()
    title(doc, TITLE, "Ett uppdateringsbart studiestöd för kursen Affärsmannaskap. Bygger på nuvarande kursmaterial, tidigare rensade dokument och pedagogiska exempel.")
    md.h1(TITLE)
    md.p("Ett uppdateringsbart studiestöd för kursen Affärsmannaskap. Bygger på nuvarande kursmaterial, tidigare rensade dokument och pedagogiska exempel.")

    h1(doc, "1. Så använder du dokumentet")
    md.h2("1. Så använder du dokumentet")
    p(doc, "Det här är ett levande instuderingsmaterial. Det betyder att dokumentet ska kunna byggas ut när kursen fortsätter och när fler föreläsningar, modeller eller exempel tillkommer.")
    bullets(doc, [
        "Täckt hittills: kursens grundlogik, konsultmässighet, erbjudande, kundvärde, analysmodeller, produkt-marknad, företagsformer, företagsplattform, mål/KPI, företagsekonomi, årsredovisning, budget/prognos, konsultbolagslönsamhet, affärsmodeller, avtalsrätt, offentlig upphandling, RFI/RFP/RFQ, affärsplan, roadmap, marknadsplan och REBR 2026 som extern källa.",
        "Kan läggas till senare: sälj- och förhandlingsteknik, presentationer, feedback, fler tentafrågor och nya case från lektioner.",
        "Studera genom att först läsa förklaringen, sedan titta på exemplet och till sist svara på kontrollfrågorna utan att titta.",
        "Miniövningarna är till för att träna användning, inte bara minne.",
        "När nytt material kommer: lägg in nya begrepp och modeller i det levande avsnittet längst bak och flytta sedan in dem i rätt del av dokumentet.",
    ])
    md.p("Det här är ett levande instuderingsmaterial. Det betyder att dokumentet ska kunna byggas ut när kursen fortsätter och när fler föreläsningar, modeller eller exempel tillkommer.")
    md.bullets([
        "Täckt hittills: kursens grundlogik, konsultmässighet, erbjudande, kundvärde, analysmodeller, produkt-marknad, företagsformer, företagsplattform, mål/KPI, företagsekonomi, årsredovisning, budget/prognos, konsultbolagslönsamhet, affärsmodeller, avtalsrätt, offentlig upphandling, RFI/RFP/RFQ, affärsplan, roadmap, marknadsplan och REBR 2026 som extern källa.",
        "Kan läggas till senare: sälj- och förhandlingsteknik, presentationer, feedback, fler tentafrågor och nya case från lektioner.",
        "Studera genom att först läsa förklaringen, sedan titta på exemplet och till sist svara på kontrollfrågorna utan att titta.",
        "Miniövningarna är till för att träna användning, inte bara minne.",
        "När nytt material kommer: lägg in nya begrepp och modeller i det levande avsnittet längst bak och flytta sedan in dem i rätt del av dokumentet.",
    ])

    h1(doc, "2. Kursen i korthet")
    md.h2("2. Kursen i korthet")
    p(doc, "Affärsmannaskap handlar om att förstå hur ett företag skapar värde för kunder och hur erbjudanden blir kommersiellt hållbara. I den här kursen kopplas det till IT, konsultrollen och projektledarens ansvar för kundnytta, ekonomi, relationer och professionellt agerande.")
    bullets(doc, [
        "För en IT-projektledare betyder affärsmannaskap att teknik måste kopplas till kundens verksamhet, nytta, risk och ekonomi.",
        "För en konsult betyder det att varje interaktion påverkar förtroende, relation och framtida affär.",
        "Kursens modeller hjälper dig gå från idé till argument: kundsegment, erbjudande, pris, marknad, konkurrenter, omvärld, bolagsform, företagsplattform, mål, ekonomisk hälsokoll, budget, affärsmodell, avtal, offentlig upphandling och affärsplan.",
        "Gruppuppgiften tränar förmågan att bygga ett fiktivt IT-konsultbolag och motivera affärslogiken steg för steg i en affärsplan.",
    ])
    p(doc, "Den röda tråden är enkel: vilka är vi, vilket problem löser vi, för vem, varför ska kunden välja oss, hur säljer och levererar vi, och hur blir det en fungerande affär?")
    md.p("Affärsmannaskap handlar om att förstå hur ett företag skapar värde för kunder och hur erbjudanden blir kommersiellt hållbara. I den här kursen kopplas det till IT, konsultrollen och projektledarens ansvar för kundnytta, ekonomi, relationer och professionellt agerande.")
    md.bullets([
        "För en IT-projektledare betyder affärsmannaskap att teknik måste kopplas till kundens verksamhet, nytta, risk och ekonomi.",
        "För en konsult betyder det att varje interaktion påverkar förtroende, relation och framtida affär.",
        "Kursens modeller hjälper dig gå från idé till argument: kundsegment, erbjudande, pris, marknad, konkurrenter, omvärld, bolagsform, företagsplattform, mål, ekonomisk hälsokoll, budget, affärsmodell, avtal, offentlig upphandling och affärsplan.",
        "Gruppuppgiften tränar förmågan att bygga ett fiktivt IT-konsultbolag och motivera affärslogiken steg för steg i en affärsplan.",
    ])
    md.p("Den röda tråden är enkel: vilka är vi, vilket problem löser vi, för vem, varför ska kunden välja oss, hur säljer och levererar vi, och hur blir det en fungerande affär?")

    add_concept_section(doc, md)
    add_models_section(doc, md)

    h1(doc, "5. Ifyllt exempel: fiktivt IT-konsultbolag")
    md.h2("5. Ifyllt exempel: fiktivt IT-konsultbolag")
    p(doc, "Pedagogiskt exempel: Verksamhetslyftet Digital AB är ett fiktivt IT-konsultbolag. Bolaget hjälper små och medelstora organisationer att minska manuell administration med AI-stöd, Microsoft 365, SharePoint, Power Automate och verksamhetsnära digitalisering.")
    md.p("Pedagogiskt exempel: Verksamhetslyftet Digital AB är ett fiktivt IT-konsultbolag. Bolaget hjälper små och medelstora organisationer att minska manuell administration med AI-stöd, Microsoft 365, SharePoint, Power Automate och verksamhetsnära digitalisering.")
    example_rows = [
        ("Målkunder", "Små och medelstora organisationer med återkommande administrativa flöden, många dokument och begränsad intern IT-kapacitet."),
        ("Kundproblem", "Manuell handläggning, dubbelarbete, otydliga dokumentytor, mejlbaserade processer och låg kontroll över uppföljning."),
        ("Värdeerbjudande", "Vi gör administration enklare och mer spårbar genom praktisk automation och AI-stöd i kundens befintliga Microsoft 365-miljö."),
        ("Tjänster", "Förstudie, SharePoint-struktur, Power Automate-flöden, enklare AI-stöd, utbildning och löpande förbättringsstöd."),
        ("Prissättning", "Fast pris för förstudie, projektpris för införande och månadsabonnemang för förbättringar och stöd."),
        ("Affärsmodell/intäktsmodell", "Kombinerar förstudier, fastprisinföranden, löpande förbättringsstöd och prenumerationsliknande support så att intäkterna inte bara bygger på engångsprojekt."),
        ("Differentiering", "Kombinerar verksamhetsförståelse med praktiskt Microsoft 365-genomförande. Fokus på små steg som kunden kan använda direkt."),
        ("Kanaler", "LinkedIn, rekommendationer, partnernätverk, seminarier för mindre organisationer och direktkontakt med verksamhetschefer."),
        ("Marknadsplan", "Bolaget prioriterar rekommendationer, LinkedIn-innehåll, korta frukostseminarier och partnerkontakter mot verksamhetschefer i dokumenttunga organisationer."),
        ("Offentlig kund/upphandling", "Om kunden är kommun, region eller myndighet behöver bolaget läsa krav, kriterier, avtal, offentlighet och överprövningsrisk innan anbud lämnas."),
        ("Intäktsströmmar", "Förstudier, införandeprojekt, utbildningar och återkommande abonnemang."),
        ("Nyckelresurser", "M365-kompetens, processkartläggning, AI-kunskap, mallar, projektledning och pedagogisk förmåga."),
        ("Nyckelpartners", "Microsoft-partner, dataskyddsjurist vid behov, redovisningsbyrå och eventuella underkonsulter."),
        ("Risker", "Kunder överskattar AI, dataskydd blir otydligt, större konsulter kopierar paketeringen eller kunden saknar tid att förändra arbetssätt."),
        ("Konkurrenter", "Lokala IT-byråer, större konsultbolag, frilansare, interna IT-avdelningar och färdiga AI-/automationsverktyg."),
        ("Möjlig tillväxtstrategi", "Börja med marknadspenetration i segmentet små och medelstora organisationer, gå sedan mot produktutveckling med standardiserade automationspaket."),
        ("Vision", "Vi vill göra smart digitalisering möjlig för organisationer som saknar stor egen IT-avdelning."),
        ("Mission", "Vi minskar administrativ friktion genom praktisk automation, AI-stöd och tydliga Microsoft 365-flöden."),
        ("Värderingar", "Tydlighet, ansvar, lärande och respekt för kundens vardag. De ska synas i statusmöten, dokumentation och rekommendationer."),
        ("Affärsidé", "Vi kartlägger, automatiserar och förbättrar administrativa processer för små och medelstora organisationer."),
        ("Mål/KPI", "Kundnöjdhet efter införande, andel leveranser i tid, återkommande intäkter och kompetenstimmar per konsult."),
        ("Budget/prognos", "Budgetera löner, sociala avgifter, pension, försäkring, licenser, utbildning, sälj och realistisk debiteringsgrad. Följ upp med prognos när beläggning eller kostnader ändras."),
        ("Ekonomisk hälsokoll", "Följ resultat, kassalikviditet, soliditet och kassaflöde så att bolaget inte bara växer på papperet utan också kan betala sina korta skulder."),
        ("Avtalsfrågor", "Klargör omfattning, ersättning, ansvar, sekretess, underleverantörer, hosting, underhåll, personuppgifter och vem som äger lösningen."),
        ("Roadmap", "Kvartal 1: förstudiepaket och första pilotkund. Kvartal 2: standardiserat införandepaket. Kvartal 3: abonnemang för förbättringsstöd. Kvartal 4: utvärdering och skalning."),
    ]
    table(doc, ["Del", "Exempel"], example_rows, [1.8, 4.7], font_size=8.5)
    md.table(["Del", "Exempel"], example_rows)

    h1(doc, "6. Företagsformer - vad jag behöver förstå")
    md.h2("6. Företagsformer - vad jag behöver förstå")
    p(doc, "Det här avsnittet bygger på kursens jämförelsematerial om företagsformer. Det är studiematerial, inte juridisk rådgivning. Poängen är att förstå hur ansvar, ägande, kapital och administration påverkar valet.")
    md.p("Det här avsnittet bygger på kursens jämförelsematerial om företagsformer. Det är studiematerial, inte juridisk rådgivning. Poängen är att förstå hur ansvar, ägande, kapital och administration påverkar valet.")
    for form, explanation, fits, advantages, disadvantages, risk_level, it_fit in COMPANY_FORMS:
        h2(doc, form)
        md.h3(form)
        rows = [
            ("Enkel förklaring", explanation),
            ("När den passar", fits),
            ("Fördelar", advantages),
            ("Nackdelar", disadvantages),
            ("Risknivå", risk_level),
            ("Passar fiktivt IT-konsultbolag?", it_fit),
        ]
        table(doc, ["Del", "Studiestöd"], rows, [1.9, 4.6], font_size=8.7)
        md.table(["Del", "Studiestöd"], rows)
    h2(doc, "Trolig rekommendation för grupparbetet")
    p(doc, "Aktiebolag är troligast att rekommendera för gruppens fiktiva IT-konsultbolag. Skälet är att det passar flera ägare, ger tydligare ansvarsskillnad mellan bolag och ägare, uppfattas som trovärdigt mot kunder och fungerar bättre om bolaget ska växa. Nackdelen är kapitalinsats och mer administration, men i ett kurscase är det lätt att motivera.")
    md.h3("Trolig rekommendation för grupparbetet")
    md.p("Aktiebolag är troligast att rekommendera för gruppens fiktiva IT-konsultbolag. Skälet är att det passar flera ägare, ger tydligare ansvarsskillnad mellan bolag och ägare, uppfattas som trovärdigt mot kunder och fungerar bättre om bolaget ska växa. Nackdelen är kapitalinsats och mer administration, men i ett kurscase är det lätt att motivera.")

    h1(doc, "7. Omvärld och REBR 2026 - hur rapporten kan användas")
    md.h2("7. Omvärld och REBR 2026 - hur rapporten kan användas")
    p(doc, "REBR 2026 är en extern källa, inte en kursmodell. Den kan användas för att göra resonemang om arbetsliv, rekrytering, arbetsgivarattraktivitet och digital arbetsmiljö mer trovärdiga.")
    md.p("REBR 2026 är en extern källa, inte en kursmodell. Den kan användas för att göra resonemang om arbetsliv, rekrytering, arbetsgivarattraktivitet och digital arbetsmiljö mer trovärdiga.")
    rebr_rows = [
        ("Från rapporten", "Balans mellan jobb och fritid, trevlig arbetsmiljö, lön/förmåner, anställningstrygghet och lika möjligheter är centrala drivkrafter vid val av arbetsgivare."),
        ("Pedagogisk tolkning", "Ett konsultbolag behöver både sälja värde till kunder och vara attraktivt för konsulter. Kompetensförsörjning blir en affärsfråga."),
        ("Exempel på tillämpning", "Verksamhetslyftet Digital AB kan argumentera för digitala arbetssätt som minskar administration och gör hybridarbete mer hanterbart."),
        ("Från rapporten", "Digitala talanger värderar flexibilitet, lättillgänglighet och praktisk vardagslogistik relativt högt."),
        ("Pedagogisk tolkning", "Employer branding för IT-konsulter bör vara konkret: arbetsbelastning, utveckling, flexibilitet och tydligt ledarskap."),
        ("Exempel på formulering", "Eftersom arbetslivsbalans och flexibilitet är viktiga drivkrafter kan vårt erbjudande kopplas till enklare digital samverkan och mindre administrativ friktion."),
        ("Varning", "Rapporten visar breda arbetsmarknadsdrivkrafter. Den bevisar inte automatiskt efterfrågan på en specifik IT-tjänst."),
    ]
    table(doc, ["Typ", "Innehåll"], rebr_rows, [1.45, 5.05], font_size=8.5)
    md.table(["Typ", "Innehåll"], rebr_rows)

    add_annual_reports_section(doc, md)
    add_lecture6_section(doc, md)
    add_lecture7_section(doc, md)

    h1(doc, "11. Instuderingsfrågor")
    md.h2("11. Instuderingsfrågor")
    q_rows = [
        ("Grundnivå", "Affärsmannaskap", "Förklara affärsmannaskap med egna ord. Hur skiljer det sig från att bara sälja?"),
        ("Grundnivå", "Kundvärde", "Vad är skillnaden mellan en funktion och kundvärde?"),
        ("Grundnivå", "FAB/5P", "Nämn delarna i FAB och 5P och ge ett kort exempel."),
        ("Grundnivå", "Företagsform", "Vad betyder juridisk person och varför spelar det roll?"),
        ("Grundnivå", "Vision/mission", "Vad är skillnaden mellan vision, mission och affärsidé?"),
        ("Grundnivå", "Resultaträkning", "Vad är skillnaden mellan nettoomsättning, rörelseresultat och årets resultat?"),
        ("Grundnivå", "Balansräkning", "Vad betyder tillgångar, skulder och eget kapital?"),
        ("Grundnivå", "Likviditet/soliditet", "Vad mäter likviditet respektive soliditet?"),
        ("Grundnivå", "Budget/prognos", "Vad är skillnaden mellan budget och prognos/forecast?"),
        ("Grundnivå", "Avtal", "Vad är skillnaden mellan ramavtal, avrop, avtal och Letter of Intent?"),
        ("Grundnivå", "Offentlig upphandling", "Vad är offentlig upphandling och varför skiljer den sig från en privat offertdialog?"),
        ("Grundnivå", "RFI/RFP/RFQ", "Vad är skillnaden mellan RFI, RFP och RFQ?"),
        ("Grundnivå", "Affärsplan/roadmap", "Vad är skillnaden mellan affärsplan, marknadsplan och roadmap?"),
        ("Tillämpning", "SWOT", "Gör en enkel SWOT för Verksamhetslyftet Digital AB."),
        ("Tillämpning", "PESTELID", "Välj tre omvärldsfaktorer som påverkar ett IT-konsultbolag och förklara påverkan."),
        ("Tillämpning", "BMC", "Fyll i fem BMC-rutor för ett bolag som säljer AI-stöd i administration."),
        ("Tillämpning", "Företagsplattform", "Skriv en vision, mission, tre värderingar och en affärsidé som hänger ihop."),
        ("Tillämpning", "Mål/KPI", "Sätt ett mål och ett KPI för kund, medarbetare, interna processer och ekonomi."),
        ("Tillämpning", "Ekonomisk hälsokoll", "Räkna enkel soliditet och likviditet för ett bolag och skriv vad siffrorna betyder."),
        ("Tillämpning", "Produkt-marknad", "Jämför två tjänster mot två kundsegment och motivera vilket alternativ som ska prioriteras."),
        ("Tillämpning", "Konsultlönsamhet", "Hur påverkas lönsamheten om debiteringsgraden sjunker eller intern tid ökar?"),
        ("Tillämpning", "Affärsmodell", "Välj två affärsmodeller från föreläsning 6 och förklara hur de skulle kunna passa ett IT-konsultbolag."),
        ("Tillämpning", "Upphandlingsanalys", "Läs en upphandling och lista vad som ska levereras, vilka villkor som gäller och vilka kriterier som avgör."),
        ("Tillämpning", "Affärsplan", "Välj fem rubriker i affärsplanen och förklara hur de hänger ihop med samma röda tråd."),
        ("Analys", "Prissättning", "När är fast pris bättre än timpris? Vilka risker finns?"),
        ("Analys", "Ansoff", "Jämför marknadspenetration och diversifiering för ett IT-konsultbolag."),
        ("Analys", "Konkurrentanalys", "Varför är interna IT-avdelningar och färdiga AI-verktyg också konkurrenter/substitut?"),
        ("Analys", "REBR/omvärld", "Hur kan REBR användas som stöd utan att överdriva slutsatserna?"),
        ("Analys", "Årsredovisning", "Hur skiljer sig Hälsö Fisk AB och Essiq AB om du jämför resultat, likviditet och soliditet?"),
        ("Analys", "Avtalsrätt", "Vilka avtalsfrågor är mest riskfyllda i ett IT-projekt: ägande, hosting, underhåll, personuppgifter, ersättning eller ansvar? Motivera."),
        ("Analys", "Offentlig upphandling", "Varför kan otydliga krav, orimliga krav eller otydlig värderingsmodell leda till överprövning?"),
        ("Analys", "Tröskelvärde/direktupphandling", "Varför är det riskabelt att plugga ett fast belopp utan att kontrollera kund, tjänst och tidpunkt?"),
    ]
    table(doc, ["Nivå", "Område", "Fråga"], q_rows, [1.05, 1.35, 4.1], font_size=8.4)
    md.table(["Nivå", "Område", "Fråga"], q_rows)

    h1(doc, "12. Miniövningar")
    md.h2("12. Miniövningar")
    ex_rows = [
        ("Skriv ett värdeerbjudande", "Skriv två meningar: kundproblem och nytta.", "Vi hjälper mindre organisationer att minska manuell administration genom AI-stöd och automation i Microsoft 365."),
        ("Välj kundsegment", "Välj ett segment och motivera med behov, betalningsvilja och åtkomst.", "små och medelstora organisationer med dokumenttunga processer, eftersom de ofta saknar intern utvecklingskapacitet men har tydliga effektivitetsproblem."),
        ("Fyll i mini-BMC", "Skriv kundsegment, värdeerbjudande, kanaler, intäkter och nyckelresurser.", "Kundsegment: små och medelstora organisationer. Värde: mindre administration. Kanal: LinkedIn och referenser. Intäkt: förstudie, projekt, abonnemang. Resurs: M365-kompetens."),
        ("Poängsätt produkt-marknad", "Jämför två tjänster mot två marknader med 1-5 på tillväxt, lönsamhet, konkurrens och kompetens.", "AI-stöd i administration till små och medelstora organisationer får hög kompetens och kundnytta; egen mjukvarutjänst till ny marknad får högre risk och lägre kompetenspoäng."),
        ("Skriv kort SWOT", "Skriv en punkt per ruta och avsluta med beslut.", "Styrka: processförståelse. Svaghet: få referenser. Möjlighet: AI-intresse. Hot: större konsulter. Beslut: börja nischat med förstudier."),
        ("Gör FAB till säljargument", "Välj en egenskap och översätt till fördel och nytta.", "SharePoint-mall ger enklare dokumentflöde och minskar tid som läggs på att leta efter rätt version."),
        ("Välj företagsform", "Välj form och motivera med ansvar, ägare och trovärdighet.", "Aktiebolag passar bäst eftersom gruppen har flera ägare och vill visa seriös konsultverksamhet med tydlig ansvarsskillnad."),
        ("Använd REBR som stöd", "Skriv ett affärsargument med källa, tolkning och avgränsning.", "REBR visar att arbetslivsbalans är viktig. Pedagogisk tolkning: digitala arbetssätt kan minska friktion. Avgränsning: rapporten bevisar inte efterfrågan på vår specifika tjänst."),
        ("Skriv företagsplattform", "Skriv vision, mission, tre värderingar och affärsidé.", "Vision: smart digitalisering för mindre organisationer. Mission: minska administration. Värderingar: tydlighet, ansvar, lärande. Affärsidé: praktisk automation i M365."),
        ("Sätt mål med BSC", "Sätt ett mål och ett KPI i varje perspektiv.", "Kund: nöjdhet efter projekt. Medarbetare: kompetenstimmar. Interna processer: leverans i tid. Finansiellt: återkommande intäkter."),
        ("Läs två årsredovisningar", "Jämför resultat, likviditet och soliditet.", "Hälsö Fisk AB visar förlust, 11 procent soliditet och svag likviditet. Essiq AB visar vinst, 28,20 procent soliditet och bättre balanslikviditet men minskat kassaläge."),
        ("Gör en enkel konsultbudget", "Lista intäkter, personalkostnader, overhead och tre scenarier för beläggning.", "Intäkt: arvoden och supportabonnemang. Kostnad: lön, sociala avgifter, pension, licenser, försäkring och sälj. Scenario: 50, 70 och 85 procent debiteringsgrad."),
        ("Välj affärsmodell", "Välj en huvudmodell och en kompletterande modell för bolaget.", "Huvudmodell: fastprisinföranden. Komplettering: prenumeration för löpande support och förbättringsstöd."),
        ("Gör avtalskarta", "Skriv fem punkter som måste regleras innan ett IT-uppdrag startar.", "Omfattning, ersättning, ägande av lösning, hosting/drift, personuppgifter, sekretess och ansvar vid försening eller ändringar."),
        ("Skilj RFI/RFP/RFQ", "Skriv ett svar på vilken typ av kundförfrågan du har framför dig.", "RFI: kunden söker marknadsinformation. RFP: kunden vill ha förslag och offert. RFQ: kunden vet vad den vill köpa och vill jämföra pris/villkor."),
        ("Läs en upphandling", "Skriv fem frågor som måste besvaras innan ni lämnar anbud.", "Vad ska levereras? För hur många användare? Vad händer med tidigare data? Vilka krav är skallkrav? Vad ger poäng i utvärderingen?"),
        ("Bygg affärsplan från modeller", "Koppla BMC, SWOT och differentiering till tre affärsplanerubriker.", "Kundsegment och värdeerbjudande från BMC går in i kunder/produkt. SWOT går in i analys. Differentiering går in i konkurrens och produkt."),
        ("Gör en roadmap", "Skriv fyra steg för första året.", "Q1: förstudiepaket och pilotkund. Q2: införandepaket. Q3: abonnemang/support. Q4: utvärdera, paketera om och skala."),
    ]
    table(doc, ["Övning", "Uppgift", "Kort exempelsvar"], ex_rows, [1.45, 2.0, 3.05], font_size=8.0)
    md.table(["Övning", "Uppgift", "Kort exempelsvar"], ex_rows)

    h1(doc, "13. Checklista inför grupparbetet")
    md.h2("13. Checklista inför grupparbetet")
    checklist = [
        "Har vi en tydlig kund, inte bara en bred marknad?",
        "Vet vi vilket kundproblem vi löser?",
        "Har vi ett konkret erbjudande som går att beskriva enkelt?",
        "Kan vi förklara kundvärdet utan tekniskt överflöd?",
        "Har vi jämfört alternativ och konkurrenter?",
        "Har vi använt modellerna som beslutsstöd, inte som dekoration?",
        "Har vi motiverat företagsform med ansvar, ägare och trovärdighet?",
        "Har vi formulerat vision, mission, värderingar och affärsidé så att de hänger ihop?",
        "Har vi mål och KPI:er för kund, medarbetare, interna processer och ekonomi?",
        "Har vi valt affärsmodell och intäktsmodell, inte bara beskrivit tjänsten?",
        "Har vi en marknadsplan: vilka kunder når vi, i vilka kanaler och med vilket budskap?",
        "Har vi budgeterat lön, sociala avgifter, pension, försäkring, licenser, sälj och realistisk debiteringsgrad?",
        "Har vi gjort en enkel ekonomisk hälsokoll med resultat, likviditet, soliditet och kassaflöde?",
        "Har vi identifierat avtalsfrågor kring scope, ersättning, ansvar, sekretess, underleverantörer, hosting, underhåll och personuppgifter?",
        "Om kunden är offentlig: har vi förstått krav, kriterier, villkor, offentlighet och överprövningsrisk?",
        "Kan vi förklara skillnaden mellan RFI, RFP och RFQ?",
        "Har vi en roadmap som visar prioriterade steg utan att bli en hel affärsplan?",
        "Täcker affärsplanen idé, kunder, marknad, konkurrens, produkt, sälj/betalning/leverans, marknadsbearbetning, affärsmodell, resurser, hållbarhet, team och analys?",
        "Har vi använt externa källor försiktigt och korrekt?",
        "Kan varje gruppmedlem förklara affärslogiken från kundproblem till intäkt?",
        "Har vi en röd tråd mellan produkt-marknad, BMC, SWOT, PESTELID, affärsplan och presentation?",
    ]
    bullets(doc, checklist)
    md.bullets(checklist)

    h1(doc, "14. Levande avsnitt för framtida kursmaterial")
    md.h2("14. Levande avsnitt för framtida kursmaterial")
    p(doc, "Använd det här avsnittet när nytt kursmaterial kommer. Skriv kort först. När något återkommer eller blir viktigt, flytta in det i rätt tidigare avsnitt.")
    md.p("Använd det här avsnittet när nytt kursmaterial kommer. Skriv kort först. När något återkommer eller blir viktigt, flytta in det i rätt tidigare avsnitt.")
    future_rows = [
        ("Nytt material att lägga till", "Anteckna föreläsning, datum och vad som verkar viktigt."),
        ("Nya begrepp", "Skriv begrepp, enkel förklaring och var i kursen det hör hemma."),
        ("Nya modeller", "Beskriv vad modellen används till och vilken fråga den svarar på."),
        ("Nya exempel", "Markera tydligt om exemplet kommer från kursmaterialet eller är ett pedagogiskt exempel."),
        ("Saker som behöver kontrolleras", "Lägg in sådant som kräver källa, förtydligande eller lärarens bekräftelse."),
        ("Frågor att ta med till lektion/grupp", "Skriv frågor som kan förbättra grupparbetets beslut och argumentation."),
    ]
    table(doc, ["Rubrik", "Hur den ska användas"], future_rows, [2.1, 4.4], font_size=8.8)
    md.table(["Rubrik", "Hur den ska användas"], future_rows)

    return doc, "\n".join(md.lines).rstrip() + "\n"


if __name__ == "__main__":
    doc, markdown = build_document()
    MD_PATH.write_text(markdown, encoding="utf-8")
    doc.save(DOCX_PATH)
    print(DOCX_PATH)
    print(MD_PATH)
