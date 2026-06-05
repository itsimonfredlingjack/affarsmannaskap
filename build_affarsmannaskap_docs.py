from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("output/docx")
OUT.mkdir(parents=True, exist_ok=True)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "F8F9FA")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    if widths:
        set_table_widths(table, widths)
    for p in doc.paragraphs[-1:]:
        p.paragraph_format.space_after = Pt(8)
    return table


def base_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(3)
    title_run = title_p.add_run(title)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(26)
    title_run.font.bold = False
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    r = sub.add_run(subtitle)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)
    return doc


def p(doc, text, style=None, bold=False):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(8 if style != "List Bullet" else 4)
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = bold
    return para


def bullets(doc, items):
    for item in items:
        p(doc, item, style="List Bullet")


def apply_numbering(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def new_decimal_num_id(doc):
    numbering = doc.part.numbering_part.element
    num = numbering.add_num(7)  # default abstract numbering for List Number.
    num.add_lvlOverride(ilvl=0).add_startOverride(1)
    return num.numId


def numbered(doc, items):
    num_id = new_decimal_num_id(doc)
    for item in items:
        para = p(doc, item, style="List Number")
        apply_numbering(para, num_id)


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def save(doc, name):
    path = OUT / name
    doc.save(path)
    return path


def build_core():
    doc = base_doc(
        "Affärsmannaskap - kärnmaterial",
        "Rensat lärmaterial från Affärsmannaskap 1-5. Fokus: förstå modellerna och använda dem i gruppens IT-konsultbolag.",
    )
    h1(doc, "1. Kursens logik")
    p(doc, "Kursen handlar inte bara om att kunna ord som kund, marknad och erbjudande. Den tränar dig i att tänka som en konsult och projektledare som måste skapa värde, ta ansvar och fatta kommersiella beslut.")
    bullets(doc, [
        "Du ska förstå affärsmannaskap och konsultmässighet i praktiska situationer.",
        "Du ska kunna identifiera, prioritera och bearbeta kunder.",
        "Du ska förstå den kommersiella delen i ett projekt: värde, ekonomi, kundnytta och långsiktiga relationer.",
        "Grupparbetet går ut på att starta ett fiktivt IT-konsultbolag och löpande redovisa delar av affären.",
        "Föreläsning 4 lägger till hur bolaget organiseras runt sin idé: vision, mission, värderingar, affärsidé, mål och uppföljning.",
        "Föreläsning 5 lägger till företagsekonomi: resultat, balansräkning, likviditet, soliditet, budget/prognos och hur årsredovisningar kan läsas.",
        "Källor: Affärsmannaskap 1, s. 18-33; Affärsmannaskap 3, s. 21; Affärsmannaskap 4, s. 3-42; AM5, s. 6-21.",
    ])
    h2(doc, "Arbetsflödet i kursen")
    numbered(doc, [
        "Börja med vilka ni är: kompetenser, ansvar, spelregler och hur gruppen ska arbeta.",
        "Välj vilken typ av IT-konsultverksamhet ni vill bygga.",
        "Formulera erbjudandet: vad säljer ni, till vem, varför köper kunden, och vilket värde skapas?",
        "Analysera marknaden: kunder, konkurrenter, omvärld, risker och möjligheter.",
        "Välj bolagsform och samarbetsform som passar affären.",
        "Formulera bolagets plattform: vision, mission, värderingar och affärsidé.",
        "Sätt mål och nyckeltal för kund, medarbetare, interna processer och ekonomi.",
        "Läs företagets ekonomi: intäkter, kostnader, resultat, tillgångar, skulder, eget kapital, likviditet och soliditet.",
        "Koppla ihop analysen till ekonomi, försäljning, presentation och tentaförståelse.",
    ])
    h1(doc, "2. Vad affärsmannaskap betyder här")
    p(doc, "I kursmaterialet kopplas affärsmannaskap ihop med relationen mellan företaget, erbjudandet, marknaden och kunden. Det är inte bara att sälja mer. Det är att skapa rätt värde på ett professionellt och långsiktigt sätt.")
    h2(doc, "Den enkla kärnfrågan")
    p(doc, "En bra sammanfattning av kursens startfråga är: vad finns det i detta för kunden? Innan ni pratar om teknik, lösning eller pris måste ni kunna svara på varför kunden ska bry sig och varför kunden ska köpa av just er.")
    add_table(doc, ["Del", "Fråga du måste kunna svara på", "Exempel i IT-konsultbolag"], [
        ["Företaget", "Vad står vi för och vilken kompetens har vi?", "Vi är specialister på säker molnflytt för mindre vård- och omsorgsaktörer."],
        ["Erbjudandet", "Vad säljer vi och hur är det paketerat?", "Kartläggning, migreringsplan, genomförande och tre månaders support."],
        ["Marknaden", "Vilka behov, konkurrenter och trender påverkar affären?", "Kunder behöver bättre säkerhet, men har begränsad intern IT-kompetens."],
        ["Kunden", "Vem betalar och vilket problem löser vi?", "VD eller verksamhetschef vill minska risk och slippa driftstopp."],
        ["Ekonomin", "Hur skapas intäkt, marginal och hållbar leverans?", "Fast pris för förstudie, timpris för införande, abonnemang för support."],
    ], [1.2, 2.4, 2.9])
    h1(doc, "3. Konsultmässighet")
    p(doc, "Konsultmässighet är hur du uppträder när du representerar både dig själv och företaget. Kursmaterialet lyfter tre hörnstenar: professionellt bemötande, kommunikativ förmåga och affärsmässighet.")
    h2(doc, "Konsultmässigt beteende")
    bullets(doc, [
        "Lyssna aktivt på kunden och anpassa kommunikationen efter mottagaren.",
        "Var lösningsorienterad: beskriv möjliga vägar framåt i stället för att bara peka på problem.",
        "Håll deadlines eller återkoppla tidigt om förutsättningar ändras.",
        "Bygg förtroende genom tydlighet, transparens och ansvarstagande.",
        "Sök feedback och förbättra leveransen över tid.",
    ])
    h2(doc, "Inte konsultmässigt")
    bullets(doc, [
        "Lova en leveranstid som ni vet är svår att hålla bara för att vinna affären.",
        "Sälja en lösning som inte passar kundens behov för att nå ett försäljningsmål.",
        "Undanhålla risker, vara otydlig, missa deadlines eller skylla ifrån sig.",
        "Glömma att varje interaktion kan påverka framtida relationer och affärer.",
        "Källa: Affärsmannaskap 1, s. 34-40.",
    ])
    h1(doc, "4. Erbjudandet")
    h2(doc, "FAB: Features, Advantages, Benefits")
    p(doc, "FAB hjälper dig att gå från vad ni gör till varför kunden bryr sig. I en presentation är det ofta nyttan, inte funktionen, som övertygar.")
    add_table(doc, ["Nivå", "Fråga", "IT-konsultexempel"], [
        ["Feature", "Vilken egenskap har tjänsten?", "Molnmigrering med säkerhetsgranskning och dokumenterad rollback-plan."],
        ["Advantage", "Vad blir bättre jämfört med alternativet?", "Mindre risk vid flytten och tydligare kontroll över ansvar."],
        ["Benefit", "Vilken nytta får kunden?", "Kunden kan byta miljö utan längre driftstopp och med bättre regelefterlevnad."],
    ], [1.2, 2.1, 3.2])
    h2(doc, "5P för erbjudandet")
    bullets(doc, [
        "Produkt: tjänsten eller lösningen ni säljer.",
        "Pris: vad kunden betalar och hur priset motiveras.",
        "Plats: var och hur kunden får tillgång till erbjudandet.",
        "Påverkan: hur kunden får veta att ni finns och varför ni är relevanta.",
        "Personal: vilka människor, kompetenser och arbetssätt som skapar förtroende.",
    ])
    h2(doc, "Pris, rabatt och tradeoff")
    p(doc, "Materialet använder principen snabbt, bra, billigt. En stark affär beskriver vilken kombination ni erbjuder och vad kunden väljer bort. Om ni lovar alla tre utan förklaring blir affären svag.")
    bullets(doc, [
        "Snabbt + bra kostar ofta mer eftersom senior kompetens och prioriterad leverans krävs.",
        "Billigt + snabbt innebär risk för lägre kvalitet, mindre dokumentation eller enklare lösning.",
        "Bra + billigt tar ofta längre tid eller kräver mer arbete från kunden.",
        "Rabatt ska inte vara standardlösningen. Motivera priset med värde, riskminskning och resultat.",
    ])
    h2(doc, "Differentiering, paketering och positionering")
    p(doc, "Differentiering betyder att ni gör er mindre jämförbara med konkurrenterna. Paketering gör köpet enklare. Positionering visar var ni vill stå i kundens huvud: pris, kvalitet, enkelhet, anpassning, hållbarhet, långsiktighet eller prestanda.")
    add_table(doc, ["Begrepp", "Användning", "Exempel"], [
        ["Differentiering", "Visa vad som skiljer er från andra.", "Nisch: säkerhetsgranskad molnflytt för mindre reglerade verksamheter."],
        ["Paketering", "Gör köpet begripligt och jämförbart.", "Bas, Standard och Trygg drift med tydligt innehåll och fast pris för förstudien."],
        ["Positionering", "Välj var ni vill ligga i marknaden.", "Inte billigast, men tryggast och enklast för kund utan egen IT-avdelning."],
    ], [1.35, 2.2, 2.95])
    h2(doc, "USP och ESP")
    p(doc, "USP handlar om det unika försäljningsargumentet: varför ert erbjudande är bättre, tydligare eller mer relevant än alternativen. ESP handlar om den känslomässiga sidan: vilken trygghet, tillit eller känsla kunden får av att välja er.")
    bullets(doc, [
        "USP-exempel: Vi garanterar dokumenterad rollback-plan och säkerhetsgenomgång i varje molnprojekt.",
        "ESP-exempel: Kunden ska känna kontroll och trygghet även om de saknar egen senior IT-kompetens.",
        "Bra affärsargument kombinerar ofta båda: konkret skillnad plus känsla av minskad risk.",
    ])
    h1(doc, "5. Analysmodellerna i rätt ordning")
    p(doc, "Modellerna är verktyg för beslut. Använd dem inte som separata skoluppgifter. Låt varje modell svara på en fråga som nästa modell behöver.")
    numbered(doc, [
        "Produkt-marknadsmatris: vilka kombinationer av kund och erbjudande verkar mest lovande?",
        "BMC: hur hänger affärsmodellen ihop?",
        "PESTLIED: vilka omvärldsfaktorer påverkar affären?",
        "Konkurrentanalys: vilka jämför kunden er med och varför?",
        "SWOT: vad betyder allt detta för våra styrkor, svagheter, möjligheter och hot?",
        "Ansoff: växer vi genom nuvarande/ny marknad och nuvarande/ny tjänst?",
        "Företagsform: vilken juridisk form passar ansvar, ägare, kapital och administration?",
        "VMV och affärsidé: vilken plattform står bolaget på och hur vill ni uppfattas?",
        "Mål/Balanced Scorecard: hur följer ni upp om bolaget rör sig åt rätt håll?",
    ])
    h2(doc, "SWOT på dig själv")
    p(doc, "En övning i materialet är att göra SWOT på sig själv. Det är inte separat från kursen: den tränar samma logik som företagsanalysen. Styrkor och svagheter är sådant du själv kan påverka. Möjligheter och hot finns i omgivningen, till exempel arbetsmarknad, teknikskiften, nätverk eller konkurrens.")
    h1(doc, "6. Plattform, organisering och mål")
    p(doc, "Föreläsning 4 flyttar grupparbetet från erbjudande och analys till hur bolaget ska stå ihop som organisation. Plattformen består av vision, mission, värderingar och affärsidé. Den ska göra det tydligt vad ni strävar mot, varför bolaget finns, hur ni beter er och vad ni gör varje dag.")
    add_table(doc, ["Del", "Fråga", "Pedagogiskt exempel"], [
        ["Vision", "Vilket framtida läge strävar bolaget mot?", "Vi vill göra smart digitalisering möjlig för organisationer utan stor IT-avdelning."],
        ["Mission", "Varför finns vi och hur bidrar vi till visionen?", "Vi förenklar administration med verksamhetsnära AI-stöd och Microsoft 365-lösningar."],
        ["Värderingar", "Vilka beteenden ska styra arbetet?", "Tydlighet, ansvar, lärande och respekt för kundens vardag."],
        ["Affärsidé", "Vad gör vi varje dag, för vem och med vilket värde?", "Vi kartlägger, automatiserar och förbättrar administrativa flöden för små och medelstora organisationer."],
    ], [1.35, 2.25, 2.9])
    h2(doc, "Mål och uppföljning")
    p(doc, "Kursmaterialet lyfter att det som mäts tenderar att bli gjort. Mål ska därför vara kopplade till hur bolaget skapar värde, inte bara till pengar. Balanced Scorecard kan användas som stöd med fyra perspektiv: kund, medarbetare, interna processer och finansiellt.")
    bullets(doc, [
        "Kundmål: till exempel kundnöjdhet, återköp eller tydlig effekt av leveransen.",
        "Medarbetarmål: till exempel beläggning, lärande, trivsel eller kompetensutveckling.",
        "Interna processer: till exempel leveransprecision, dokumentation, kvalitet och förbättringstakt.",
        "Finansiella mål: till exempel intäkt, marginal, kassaflöde eller andel återkommande intäkter.",
        "Delmål behövs så att gruppen kan stanna upp, följa upp och justera innan slutpresentationen.",
    ])
    h1(doc, "7. Företagsekonomi och årsredovisning")
    p(doc, "Föreläsning 5 flyttar ekonomidelen från allmänna ord till sådant som går att läsa i ett bolag: resultaträkning, balansräkning, årsredovisning och nyckeltal. Poängen är inte att bli revisor, utan att kunna se om en affär verkar lönsam, betalningsförmögen och långsiktigt hållbar.")
    add_table(doc, ["Begrepp", "Vad det visar", "Fråga att ställa"], [
        ["Resultaträkning", "Intäkter minus kostnader under en period. Visar om verksamheten går med vinst eller förlust.", "Tjänar bolaget pengar på sin affär, eller äter kostnaderna upp intäkterna?"],
        ["Balansräkning", "Tillgångar, skulder och eget kapital vid en viss tidpunkt.", "Hur är bolaget finansierat, och hur mycket risk ligger i skulder?"],
        ["Eget kapital", "Ägarnas kapital i bolaget: tillgångar minus skulder. I aktiebolag skiljer man mellan bundet och fritt eget kapital.", "Finns det buffert, eller är bolaget nära att äta upp sitt kapital?"],
        ["Likviditet", "Kortsiktig betalningsförmåga. Kassalikviditet jämför omsättningstillgångar exklusive lager med kortfristiga skulder.", "Kan bolaget betala sina korta skulder utan att allt bygger på framtida försäljning?"],
        ["Soliditet", "Eget kapital i relation till totala tillgångar. Visar finansiell motståndskraft.", "Är bolaget finansierat med egen styrka eller hårt beroende av skulder?"],
        ["Kassaflöde", "Pengar in och ut över tid. Ett bolag kan vara lönsamt men ändå få likviditetsproblem.", "Kommer pengarna in i tid för att betala löner, leverantörer och skatt?"],
    ], [1.35, 3.05, 2.1])
    h2(doc, "Så läser du en årsredovisning snabbt")
    numbered(doc, [
        "Börja med förvaltningsberättelsen: vad säger bolaget själv har hänt under året?",
        "Titta på resultaträkningen: nettoomsättning, rörelseresultat, resultat efter finansiella poster och årets resultat.",
        "Titta på balansräkningen: tillgångar, eget kapital, kortfristiga skulder och likvida medel.",
        "Räkna eller läs nyckeltal: likviditet och soliditet räcker långt för en första hälsokoll.",
        "Läs revisionsberättelsen om den finns: finns anmärkningar eller tillstyrker revisorn resultat- och balansräkningen?",
    ])
    h2(doc, "Varningssignaler")
    bullets(doc, [
        "Negativt resultat kan vara hanterbart ett år, men blir allvarligt om det kombineras med låg soliditet och svag likviditet.",
        "Hög omsättning räcker inte om marginalen är svag eller kassaflödet pressat.",
        "Likviditet under 100 procent betyder att kortfristiga skulder är större än omsättningstillgångarna. Det kräver förklaring.",
        "Soliditet under cirka 30-40 procent kan vara en svaghetsmarkör enligt AM5, men bransch och bolagsfas spelar roll.",
    ])
    h1(doc, "8. Checklista inför grupparbetet")
    bullets(doc, [
        "Kan ni beskriva erbjudandet i en mening utan fackspråk?",
        "Har ni valt kundsegment och förklarat varför just det segmentet har behovet?",
        "Har ni jämfört minst tre konkurrenter eller substitut?",
        "Har ni motiverat pris och paketering med kundnytta, inte bara kostnad?",
        "Har ni använt extern data där det stärker marknadsresonemanget?",
        "Har ni valt bolagsform med tydligt skäl kopplat till ansvar, ägare och trovärdighet?",
        "Har ni formulerat vision, mission, värderingar och affärsidé så att de hänger ihop?",
        "Har ni mål och KPI:er för kund, medarbetare, interna processer och ekonomi?",
        "Har ni gjort en enkel ekonomisk hälsokoll: intäkt, kostnad, resultat, likviditet, soliditet och kassaflöde?",
        "Inför delredovisning 10 juni 2026: gör företagsform, erbjudande, paketering, differentiering, positionering, konkurrentanalys, SWOT och PEST/PESTLIED. Gör INTE mer än det som efterfrågas.",
        "Har ni en röd tråd från analys till beslut?",
    ])
    return save(doc, "01_affarsmannaskap_karnmaterial.docx")


def build_workbook():
    doc = base_doc(
        "Modellverkstad - arbetsexempel för IT-konsultbolag",
        "Praktiskt stöd för BMC, PESTLIED, konkurrentanalys, SWOT, Ansoff, produkt-marknad, företagsform, företagsplattform, mål och ekonomi.",
    )
    h1(doc, "Fiktivt exempel som används i dokumentet")
    p(doc, "Exemplet nedan använder ett påhittat bolag: SecureFlow IT AB. Bolaget hjälper små och medelstora organisationer att flytta till molnet med fokus på säkerhet, dokumentation och driftsäkerhet.")
    h1(doc, "1. Produkt-marknadsmatris")
    p(doc, "Matrisen används för att välja vilka erbjudanden och kundsegment som är mest värda att prioritera. Den ska inte fyllas i för att få fina färger, utan för att tvinga fram ett beslut.")
    p(doc, "Originalmallen är bred och uppdelad över två PDF-sidor. I praktiken betyder det att ni kan jämföra flera produkter/tjänster mot flera marknader/kundgrupper, men ni behöver inte använda fler rutor än ni faktiskt kan resonera om.")
    h2(doc, "Så använder du den")
    numbered(doc, [
        "Skriv 3-5 möjliga produkter eller tjänster överst.",
        "Skriv 3-5 möjliga kundsegment längs sidan.",
        "Ge varje kombination 1-5 poäng på tillväxt, lönsamhet, konkurrens och kompetens.",
        "Räkna ihop och markera de kombinationer som bör prioriteras.",
        "Kontrollera att hög poäng inte bara betyder att idén känns rolig. Den ska gå att sälja och leverera.",
    ])
    add_table(doc, ["Kundsegment", "Tjänst", "Tillväxt", "Lönsamhet", "Konkurrens", "Kompetens", "Summa", "Beslut"], [
        ["Privata vårdgivare", "Säker molnflytt", "4", "4", "3", "5", "16", "Prioritera"],
        ["Lokala butiker", "Webbshopspaket", "3", "3", "2", "3", "11", "Vänta"],
        ["Kommunala bolag", "IT-strategi", "3", "4", "2", "4", "13", "Utred mer"],
    ], [1.35, 1.35, 0.65, 0.75, 0.85, 0.75, 0.6, 1.2])
    h1(doc, "2. Business Model Canvas")
    p(doc, "BMC visar hur affären hänger ihop. För kursen är den viktig eftersom den kopplar erbjudandet till kunder, kanaler, resurser, aktiviteter, kostnader och intäkter.")
    add_table(doc, ["BMC-ruta", "Fråga", "SecureFlow-exempel"], [
        ["1. Kundsegment", "Vem skapar vi värde för?", "Små och medelstora verksamheter med känslig data men begränsad intern IT."],
        ["2. Värdeerbjudande", "Vilket problem löser vi?", "Trygg molnflytt med minskad driftstörning, tydlig dokumentation och säkerhetskontroll."],
        ["3. Kundrelationer", "Hur bygger vi förtroende?", "Förstudie, workshop, veckovis status, uppföljning efter införande."],
        ["4. Kanaler", "Hur når och levererar vi?", "LinkedIn, partnernätverk, rekommendationer, direktkontakt, digitala möten."],
        ["5. Partners", "Vilka behöver vi?", "Molnleverantör, säkerhetspartner, jurist/dataskydd vid behov."],
        ["6. Nyckelaktiviteter", "Vad måste vi göra bra?", "Kartläggning, migration, test, dokumentation, utbildning."],
        ["7. Nyckelresurser", "Vad måste vi ha?", "Senior molnkompetens, projektledning, säkerhetsmallar, referenscase."],
        ["8. Kostnadsstruktur", "Vad kostar affären?", "Löner, certifieringar, verktyg, försäkring, säljtid, underkonsulter."],
        ["9. Intäktsströmmar", "Hur tjänar vi pengar?", "Fast pris för förstudie, projektarvode, supportabonnemang."],
        ["10. Påverkan/cirkularitet", "Hur kan hållbarhet in i modellen?", "Mindre lokal hårdvara, längre livslängd på utrustning, resfria leveranser när det passar."],
    ], [1.55, 2.05, 2.9])
    h1(doc, "3. PESTLIED")
    p(doc, "PESTLIED hjälper dig att se omvärldsfaktorer som kan påverka affären. Varje faktor bör bedömas med påverkan, tidsram, typ och riktning. Det viktiga är inte att fylla i allt, utan att hitta det som faktiskt påverkar ert bolag.")
    bullets(doc, [
        "H/M/L/O betyder hög, mellan, låg eller osäker potentiell inverkan.",
        "Tidsramen kan exempelvis vara 6-12 månader, 12-24 månader eller längre än 24 månader.",
        "Typen visar om faktorn främst är positiv eller negativ.",
        "Riktningen visar om påverkan ökar, är oförändrad eller minskar.",
    ])
    add_table(doc, ["Faktor", "Exempel", "Bedömning"], [
        ["Politisk/legal", "Ökade krav på informationssäkerhet och dokumentation.", "Hög påverkan, 6-12 månader, positiv om vi kan visa kompetens."],
        ["Ekonomisk", "Kunder pressas av kostnader och vill undvika stora engångsprojekt.", "Mellan påverkan, 12-24 månader, kräver tydlig ROI och paketering."],
        ["Social", "Fler vill ha flexibelt arbete och stabil IT-miljö.", "Mellan påverkan, positiv, stärker moln- och samarbetslösningar."],
        ["Teknologisk", "AI, moln och säkerhetsautomation förändrar leveransen.", "Hög påverkan, ökande, kräver kompetensutveckling."],
        ["Internationell", "Leverantörer och molntjänster styrs av globala aktörer.", "Mellan påverkan, osäker, kräver leverantörsval och riskbedömning."],
        ["Miljö", "Mindre lokal hårdvara och färre resor kan vara säljargument.", "Låg till mellan påverkan, positiv om kunden värderar hållbarhet."],
        ["Demografisk", "Brist på intern IT-kompetens hos mindre organisationer.", "Hög påverkan, positiv för konsultbehov."],
    ], [1.25, 3.0, 2.25])
    doc.add_page_break()
    h1(doc, "4. Konkurrentanalys")
    p(doc, "Konkurrenter är inte bara bolag som gör exakt samma sak. Det kan också vara interna IT-avdelningar, frilansare, större konsultbolag, produktleverantörer eller att kunden inte gör något alls.")
    add_table(doc, ["Fråga", "Vad du ska leta efter", "Exempel"], [
        ["Vilka är de?", "Direkt, indirekt eller substitut.", "Stor konsultfirma, frilanskonsult, molnleverantörens egna tjänster."],
        ["Vad erbjuder de?", "Tjänst, produkt, abonnemang, paket.", "Migration, driftavtal, säkerhetsgranskning."],
        ["Hur prissätter de?", "Fast pris, timpris, paket, volymrabatt.", "Lågt timpris men otydlig dokumentation."],
        ["Vilka kunder tar de?", "Storlek, bransch, nisch.", "Stora kunder med ramavtal eller små kunder via rekommendationer."],
        ["Styrkor/svagheter", "Varför väljer eller väljer bort kunden dem?", "Starkt varumärke men dyrt och mindre personligt."],
    ], [1.5, 2.3, 2.7])
    h1(doc, "5. SWOT")
    p(doc, "SWOT ska sammanfatta analysen, inte ersätta den. Styrkor och svagheter är interna. Möjligheter och hot är externa.")
    add_table(doc, ["Ruta", "SecureFlow-exempel", "Vad beslutet blir"], [
        ["Styrka", "Hög kompetens inom moln, säkerhet och projektledning.", "Sälj trygghet och kvalitet, inte lägsta pris."],
        ["Svaghet", "Nytt bolag utan många referenser.", "Skapa pilotcase, partnerbevis och tydlig process."],
        ["Möjlighet", "Kunder behöver säker digitalisering men saknar egen kompetens.", "Välj segment där behovet är akut och budget finns."],
        ["Hot", "Stora konsultbolag kan ta ramavtal och pressa synlighet.", "Nischa erbjudandet och använd partnerkanaler."],
    ], [1.2, 3.0, 2.3])
    doc.add_page_break()
    h1(doc, "6. Ansoff")
    p(doc, "Ansoff-matrisen hjälper er att beskriva tillväxtrisk. Lägre risk betyder inte alltid bättre, men högre risk kräver bättre bevis.")
    add_table(doc, ["Strategi", "Vad det betyder", "Exempel", "Risk"], [
        ["Marknadspenetration", "Sälj mer av nuvarande tjänst till nuvarande marknad.", "Fler molnflyttar till samma kundsegment.", "Låg"],
        ["Produktutveckling", "Ny tjänst till befintliga kunder.", "Lägg till säkerhetsövervakning efter migration.", "Medium"],
        ["Marknadsutveckling", "Samma tjänst till ny kundtyp.", "Gå från privata vårdgivare till skolor.", "Medium"],
        ["Diversifiering", "Ny tjänst till ny marknad.", "AI-rådgivning för industribolag utan tidigare relation.", "Hög"],
    ], [1.55, 2.2, 2.25, 0.8])
    h1(doc, "7. Företagsform")
    p(doc, "Företagsformen ska väljas utifrån ansvar, antal ägare, kapital, trovärdighet och administration. Företagsformer 2025 ska användas som kursens jämförelseunderlag, inte som juridisk rådgivning. För ett gruppcase i IT-konsultbranschen är aktiebolag ofta lättast att motivera eftersom bolaget är juridisk person och ägarnas privata ansvar normalt är avskilt från bolagets skulder.")
    add_table(doc, ["Form", "När den kan passa", "Viktig begränsning"], [
        ["Enskild näringsverksamhet", "En person testar en liten verksamhet med låg risk.", "Personligt ansvar för företagets skulder."],
        ["Aktiebolag", "Flera ägare, konsultaffärer, behov av trovärdighet och tydlig ansvarsskillnad.", "Kapitalkrav minst 25 000 kr och mer administration."],
        ["Handelsbolag", "Minst två personer driver enklare verksamhet tillsammans.", "Bolagsmännen är personligt och solidariskt ansvariga."],
        ["Kommanditbolag", "En ansvarig part och andra med begränsat ansvar.", "Komplementären har personligt ansvar."],
        ["Ekonomisk förening", "Minst tre medlemmar samverkar för medlemmarnas ekonomiska nytta.", "Passar sämre om caset är ett vanligt konsultbolag med aktieägare."],
    ], [1.45, 3.1, 2.0])
    h1(doc, "8. Företagsplattform: vision, mission, värderingar och affärsidé")
    p(doc, "Föreläsning 4 visar att bolaget behöver en gemensam plattform. Plattformen ska inte vara fluffig text vid sidan av affären. Den ska hjälpa gruppen fatta beslut, prioritera och förklara varför bolaget finns.")
    add_table(doc, ["Del", "Vad ni ska skriva", "SecureFlow-exempel"], [
        ["Vision", "En riktning som skapar motivation och visar vad ni strävar mot.", "Att göra säker digitalisering enkel för organisationer utan stor IT-avdelning."],
        ["Mission", "Varför ni finns och hur ni bidrar till visionen.", "Vi hjälper kunder att flytta och förbättra digitala arbetssätt utan att tappa kontroll, säkerhet eller vardagsnytta."],
        ["Värderingar", "De beteenden som ska styra ledare och medarbetare.", "Tydlighet, ansvarstagande, lärande och respekt för kundens verklighet."],
        ["Affärsidé", "Vad ni gör varje dag, för vem och med vilket värde.", "Vi kartlägger, säkrar och förbättrar moln- och samarbetsmiljöer för mindre verksamheter med höga krav på ordning och trygghet."],
    ], [1.25, 2.45, 2.8])
    h1(doc, "9. Mål, KPI och Balanced Scorecard")
    p(doc, "Mål ska visa vart bolaget ska och hur gruppen vet att det går åt rätt håll. Kursmaterialet lyfter Balanced Scorecard som möjlig modell: kund, medarbetare, interna processer och finansiella mål.")
    add_table(doc, ["Perspektiv", "Möjligt mål", "Möjligt KPI"], [
        ["Kund", "Trygg leverans och tydlig nytta.", "Kundnöjdhet, återkommande kunder, upplevd effekt."],
        ["Medarbetare", "Konsulter utvecklas och vill stanna.", "Kompetenstimmar, trivsel, hållbar beläggning."],
        ["Interna processer", "Spårbar leverans och jämn kvalitet.", "Leverans i tid, dokumenterad process, förbättringar."],
        ["Finansiellt", "Lönsamhet och återkommande intäkter.", "Marginal, kassaflöde, intäkt per konsult, abonnemangsandel."],
    ], [1.35, 3.1, 2.05])
    h1(doc, "10. Ekonomisk hälsokoll")
    p(doc, "AM5 visar hur gruppen kan gå från affärsidé till ekonomisk kontroll. I ett fiktivt IT-konsultbolag räcker det ofta med en enkel minitabell: vad säljer vi för, vad kostar leveransen, hur snabbt får vi betalt, och hur mycket buffert har bolaget?")
    add_table(doc, ["Nyckel", "SecureFlow-exempel", "Tolkning"], [
        ["Resultat", "Nettoomsättning 2 400 000 kr, kostnader 2 050 000 kr, resultat 350 000 kr.", "Affären verkar lönsam, men marginalen måste räcka till risk, säljtid och utveckling."],
        ["Täckningsbidrag", "Ett molnprojekt säljs för 180 000 kr och direkta konsult-/licenskostnader är 105 000 kr. TB blir 75 000 kr.", "Projektet bidrar till fasta kostnader som lön, administration och sälj."],
        ["Kassalikviditet", "Omsättningstillgångar exkl. lager 620 000 kr / kortfristiga skulder 410 000 kr = 151 procent.", "Över 100 procent tyder på att korta skulder kan betalas utan akut stress."],
        ["Soliditet", "Eget kapital 360 000 kr / totala tillgångar 900 000 kr = 40 procent.", "Bolaget har rimlig motståndskraft och är inte enbart skuldfinansierat."],
        ["Kassaflöde", "Tre stora kunder betalar efter 45 dagar medan löner betalas varje månad.", "Lönsamhet räcker inte om pengarna kommer in för sent. Likviditetsbudget behövs."],
    ], [1.25, 3.2, 2.05])
    h2(doc, "Skriv så här i grupparbetet")
    bullets(doc, [
        "Vi bedömer inte bara om idén går att sälja, utan också om bolaget klarar vardagen ekonomiskt.",
        "Våra viktigaste finansiella nyckeltal är kassalikviditet, soliditet, resultat och andel återkommande intäkter.",
        "Om likviditeten blir svag prioriterar vi kortare betalningstid, mindre förskottskostnader eller abonnemang.",
        "Om soliditeten blir svag behöver vi bygga eget kapital innan vi tar stora risker.",
    ])
    h1(doc, "11. Skrivmall för er egen analys")
    numbered(doc, [
        "Vårt valda kundsegment är ... därför att ...",
        "Vårt erbjudande löser problemet ... genom ...",
        "Kunden får värdet ... vilket kan mätas genom ...",
        "Våra viktigaste konkurrenter/substitut är ...",
        "Vår differentiering är ...",
        "Våra största risker är ... och vi hanterar dem genom ...",
        "Den bolagsform vi väljer är ... eftersom ...",
        "Vår vision är ... och den skapar riktning genom ...",
        "Vår mission är ... och vårt bidrag till kunden är ...",
        "Våra värderingar är ... och de ska synas i beteenden som ...",
        "Våra viktigaste mål/KPI:er är ... därför att ...",
        "Vår ekonomiska hälsokoll visar ... eftersom resultat, likviditet och soliditet ...",
    ])
    return save(doc, "02_modellverkstad_arbetsexempel.docx")


def build_rebr():
    doc = base_doc(
        "REBR 2026 som extern källa",
        "Kort stöd för hur Randstad Employer Brand Research 2026 kan användas i Affärsmannaskap utan att överdriva slutsatserna.",
    )
    h1(doc, "Varför rapporten är användbar")
    p(doc, "REBR 2026 är inte en kursmodell. Den är en extern källa som kan ge stöd när ni diskuterar marknad, rekrytering, employer value proposition, positionering och behov hos kunder eller kandidater. Använd den som belägg, inte som facit.")
    h1(doc, "Kursrelevanta datapunkter")
    bullets(doc, [
        "Balans mellan jobb och fritid samt trevlig arbetsmiljö anges som de främsta anledningarna till att talanger i Sverige väljer arbetsgivare. Källa: REBR 2026, s. 4 och 9.",
        "Lön och förmåner är fortfarande mycket viktiga och rankas högt när människor tvingas prioritera mellan faktorer. Källa: REBR 2026, s. 9-12.",
        "Anställningstrygghet och lika möjligheter finns bland de fem viktigaste drivkrafterna. Källa: REBR 2026, s. 9.",
        "Rapporten beskriver ett tydligt gap mellan ideal och upplevelse för lön/förmåner samt balans mellan jobb och fritid. Källa: REBR 2026, s. 13.",
        "För känslan av anställningstrygghet lyfts pålitlig lön/förmåner, rättvisa arbetsmetoder och transparent kommunikation. Källa: REBR 2026, s. 15.",
        "För balans mellan jobb och fritid lyfts rimlig arbetsbelastning, återhämtning och bra arbetsmiljö. Källa: REBR 2026, s. 17.",
        "Cirka 35 procent arbetar på distans åtminstone en del av tiden, enligt rapportens distansarbetsavsnitt. Källa: REBR 2026, s. 18.",
        "22 procent planerar att byta arbetsgivare under första halvåret 2026, medan 11 procent hade bytt arbetsgivare under de sista sex månaderna 2025. Källa: REBR 2026, s. 20.",
        "Främsta skälen att lämna arbetsgivare är för låg ersättning, dålig relation till ledningen, bättre balans mellan jobb och fritid samt bristande intresse för jobbet. Källa: REBR 2026, s. 21-22.",
        "Digitala talanger lägger relativt stor vikt vid bekvämlighet, lättillgänglighet och flexibla arbetsformer. Källa: REBR 2026, s. 7.",
        "Arbetsförmedlingen, LinkedIn, personliga kontakter och jobbportaler behandlas som viktiga jobbsökningskanaler. Källa: REBR 2026, s. 23-24.",
        "Många arbetssökande värdesätter personlig kontakt även i en digital jobbsökningsprocess. Källa: REBR 2026, s. 25.",
    ])
    h1(doc, "Så kan rapporten stärka ert IT-konsultcase")
    add_table(doc, ["Del i affärscaset", "Hur REBR kan användas", "Exempelmening"], [
        ["Kundbehov", "Visa att arbetsmiljö, flexibilitet och trygghet är viktiga arbetslivsfrågor.", "Vår tjänst stärker distansarbete och minskar friktion i vardagen."],
        ["Erbjudande", "Koppla IT-lösningen till konkreta effekter för medarbetare och ledning.", "Säker digital samverkan gör hybridarbete enklare utan att tumma på kontroll."],
        ["Positionering", "Differentiera er mot ren teknikförsäljning genom arbetslivsnytta.", "Vi säljer inte bara moln, vi säljer tryggare digital arbetsmiljö."],
        ["Rekrytering", "Visa hur bolaget själv kan locka konsulter.", "Vi behöver erbjuda tydlig utveckling, rimlig belastning och transparent ledarskap."],
        ["Riskanalys", "Använd rapportens jobbytesorsaker som hot i SWOT.", "Dåligt ledarskap eller svag utveckling kan göra det svårt att behålla konsulter."],
    ], [1.55, 2.45, 2.5])
    h1(doc, "Exempel: från källa till affärsargument")
    h2(doc, "Svagt argument")
    p(doc, "Alla vill jobba hemifrån, därför ska vi sälja distanslösningar.")
    h2(doc, "Starkare argument")
    p(doc, "REBR 2026 visar att balans mellan jobb och fritid är en central arbetsgivardrivkraft och att distansarbete är etablerat för en del av arbetsmarknaden. Därför kan ett IT-konsultbolag motivera tjänster som gör hybridarbete säkrare, tydligare och lättare att använda. Argumentet bör avgränsas till kunder där sådana arbetsformer faktiskt är möjliga.")
    h1(doc, "Vad ni inte ska överdriva")
    bullets(doc, [
        "Säg inte att rapporten bevisar efterfrågan på just er tjänst. Den visar breda arbetsmarknadsdrivkrafter.",
        "Säg inte att alla yrkesgrupper vill samma sak. Rapporten visar skillnader mellan generationer, kön och yrkeskategorier.",
        "Säg inte att distansarbete är relevant för alla kunder. Vissa roller kan inte utföras på distans.",
        "Blanda inte ihop kandidaters arbetsgivarpreferenser med kunders köpvilja. De kan kopplas ihop, men är inte samma sak.",
        "Använd procenttal med källa och sammanhang. Undvik att stapla siffror utan att förklara affärsbetydelsen.",
    ])
    h1(doc, "Kort källtext att använda i grupparbetet")
    p(doc, "Randstad Employer Brand Research 2026 visar att svenska talanger prioriterar balans mellan jobb och fritid, trevlig arbetsmiljö, lön och förmåner samt anställningstrygghet vid val av arbetsgivare. För ett IT-konsultbolag kan detta användas som stöd för erbjudanden som förbättrar digital arbetsmiljö, hybridarbete, säkerhet och ledningens förmåga att skapa tydliga arbetsformer. Slutsatsen bör avgränsas: rapporten visar arbetsmarknadsdrivkrafter, inte automatisk efterfrågan på en specifik IT-tjänst.")
    h1(doc, "Källnot")
    p(doc, "Källa: Randstad Employer Brand Research 2026, Country Report Sweden, maj 2026. Urval enligt rapporten: Sverige, 3895 respondenter; fältarbete januari 2026; intervjuer av Kantar. Källa: REBR 2026, s. 32.")
    return save(doc, "03_rebr_2026_extern_kalla.docx")


if __name__ == "__main__":
    paths = [build_core(), build_workbook(), build_rebr()]
    for path in paths:
        print(path)
